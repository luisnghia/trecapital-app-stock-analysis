from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
import html
import traceback
import math
import textwrap
from typing import Any, Iterable

import pandas as pd
import plotly.graph_objects as go

from module1_engine import (
    CompanyOverview,
    latest_metric_cards,
    build_flags,
    build_quick_summary,
    build_metric_dict,
    build_fcf_analysis_table,
    build_cashflow_scorecard,
    build_cashflow_situation_alerts,
    build_financial_ratio_table,
    build_financial_ratio_scorecard,
    build_financial_ratio_alerts,
    build_value_investing_assessment,
    build_mos_valuation_table,
    build_mos_detailed_summary,
    build_combined_assessment_table,
    format_table_for_display,
)
from module1_charts import (
    make_line_fig,
    make_dupont_profitability_fig,
    make_dupont_driver_fig,
    make_roic_investment_fig,
    make_fcf_generation_fig,
    make_fcf_usage_fig,
    make_fcf_conversion_fig,
)
from module2_engine import (
    load_assumptions,
    classify_company,
    build_module2_valuation_table,
    build_valuation_range,
    build_porter_moat_scorecard,
    build_value_chain_table,
    build_risk_scenario_table,
    build_module2_summary,
    build_beneish_mscore_table,
    build_accrual_quality_table,
    build_modified_jones_kothari_table,
    build_real_earnings_management_table,
)

APP_DIR = Path(__file__).resolve().parent
DEFAULT_ASSUMPTIONS_PATH = APP_DIR / "configs" / "valuation_assumptions.json"
DEFAULT_REPORT_DIR = APP_DIR / "reports"

EXPORT_FORMATS = {
    "Excel (.xlsx)": ("xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    "Word (.docx)": ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "PDF (.pdf)": ("pdf", "application/pdf"),
}


PRINT_PAGE_CSS = r"""
@media print {
  @page { size: A4 landscape; margin: 8mm; }
  html, body { background: #FFFFFF !important; }
  body { -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }
  [data-testid="stSidebar"], [data-testid="stHeader"], [data-testid="stToolbar"], [data-testid="stDecoration"], #MainMenu, footer { display: none !important; }
  .stDeployButton, .stButton, .stDownloadButton, .stSelectbox, .stCheckbox, .stTextInput, .stSlider, .stExpander [data-testid="stExpanderToggleIcon"] { display: none !important; }
  [data-testid="stAppViewContainer"] { margin: 0 !important; padding: 0 !important; }
  [data-testid="stMain"] { margin: 0 !important; padding: 0 !important; }
  .block-container { max-width: none !important; width: 100% !important; padding: 0 !important; margin: 0 !important; }
  section.main > div { padding: 0 !important; }
  div[data-testid="stExpander"] { border: 0 !important; box-shadow: none !important; break-inside: avoid; page-break-inside: avoid; }
  div[data-testid="stVerticalBlock"], div[data-testid="stHorizontalBlock"], .element-container { break-inside: avoid; page-break-inside: avoid; }
  iframe { max-width: 100% !important; border: 0 !important; }
  .page-brand-shell, .hero-card, .trecapital-card, .note-card, .metric-card, .reason-card, .price-card { break-inside: avoid; page-break-inside: avoid; }
  .stTabs [role="tablist"] { display: none !important; }
}
"""


def _inject_print_page_css(st_module: Any) -> None:
    """Inject browser print CSS so PDF output follows the visible Streamlit app layout."""
    try:
        st_module.markdown(f"<style id='trecapital-print-page-css'>{PRINT_PAGE_CSS}</style>", unsafe_allow_html=True)
    except Exception:
        pass


def _render_browser_print_pdf_button(
    widget_key: str = "full_report_export",
    *,
    button_text: str = "🖨️ In trang hiện tại / Save as PDF - giữ nguyên format app",
    help_text: str | None = None,
) -> None:
    """Render a browser-side print button.

    Streamlit's Python process cannot see the fully rendered browser DOM, so the most reliable way
    to keep the dashboard's exact layout/colors/cards is to call the browser print dialog and let the
    user choose "Save as PDF". V23.37 also uses this button on the consolidated report page, which
    renders every module in one long app-like report before printing.
    """
    import streamlit.components.v1 as components
    css_js = PRINT_PAGE_CSS.replace("`", "\\`")
    if help_text is None:
        help_text = "Khi hộp thoại in mở ra, chọn máy in <b>Save as PDF</b> hoặc <b>Microsoft Print to PDF</b>, khổ giấy <b>A4 ngang</b>, bật <b>Background graphics</b> để giữ màu card/bảng."
    components.html(
        f"""
        <div style="font-family:system-ui,-apple-system,Segoe UI,sans-serif; padding:10px 2px 4px 2px;">
          <button id="printPdfBtn_{html.escape(widget_key)}" style="
              width:100%; border:1.5px solid rgba(11,127,117,.35); border-radius:16px;
              background:linear-gradient(135deg,#FFF3C4 0%,#FFFFFF 100%); color:#064E47;
              font-size:15px; font-weight:900; padding:13px 16px; cursor:pointer;
              box-shadow:0 8px 22px rgba(11,127,117,.12);">
            {html.escape(button_text)}
          </button>
          <div style="font-size:12px; color:#64748B; margin-top:8px; line-height:1.45;">
            {help_text}
          </div>
        </div>
        <script>
        (function() {{
          const btn = document.getElementById('printPdfBtn_{html.escape(widget_key)}');
          function installCss() {{
            try {{
              const parentDoc = window.parent.document;
              let style = parentDoc.getElementById('trecapital-print-page-css-js');
              if (!style) {{
                style = parentDoc.createElement('style');
                style.id = 'trecapital-print-page-css-js';
                style.innerHTML = `{css_js}`;
                parentDoc.head.appendChild(style);
              }}
            }} catch (e) {{ /* Parent CSS is also injected by Streamlit markdown. */ }}
          }}
          if (btn) {{
            btn.addEventListener('click', function() {{
              installCss();
              setTimeout(function() {{
                try {{ window.parent.focus(); window.parent.print(); }}
                catch(e) {{ window.print(); }}
              }}, 180);
            }});
          }}
          installCss();
        }})();
        </script>
        """,
        height=116,
        scrolling=False,
    )

# Trecapital visual identity used by the app and by exported reports.
BRAND_TEAL = "0B7F75"
BRAND_TEAL_DARK = "064E47"
BRAND_TEAL_SOFT = "EAF7F1"
BRAND_YELLOW = "F5B21B"
BRAND_YELLOW_SOFT = "FFF3C4"
BRAND_RED = "DC2626"
BRAND_RED_SOFT = "FEE2E2"
BRAND_GREEN_SOFT = "D1FAE5"
BRAND_GREEN_DARK = "064E3B"
BRAND_BORDER = "CBD5E1"



@dataclass
class ReportSection:
    title: str
    paragraphs: list[str]
    tables: list[tuple[str, pd.DataFrame]]
    figures: list[tuple[str, go.Figure]]


@dataclass
class ReportPackage:
    title: str
    ticker: str
    sections: list[ReportSection]
    warnings: list[str]


def _safe_ticker(value: object) -> str:
    text = str(value or "").upper().strip()
    text = re.sub(r"[^A-Z0-9_.-]+", "_", text)
    return text or "REPORT"


def _safe_sheet_name(value: str, used: set[str]) -> str:
    name = re.sub(r"[\\/*?:\[\]]", " ", str(value or "Sheet")).strip()[:31] or "Sheet"
    base = name
    idx = 2
    while name in used:
        suffix = f" {idx}"
        name = (base[:31 - len(suffix)] + suffix).strip()
        idx += 1
    used.add(name)
    return name


def _clean_text(value: object) -> str:
    text = "" if value is None else str(value)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    text = text.replace("**", "").replace("__", "")
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    return text


def _df(df: Any) -> pd.DataFrame:
    if df is None:
        return pd.DataFrame()
    if isinstance(df, pd.DataFrame):
        return df.copy()
    try:
        return pd.DataFrame(df)
    except Exception:
        return pd.DataFrame()


def _fig_to_png_bytes(fig: go.Figure | None, width: int = 1200, height: int = 720, scale: int = 2) -> bytes | None:
    """Render Plotly figure to PNG.

    Primary path uses Plotly/Kaleido. If Kaleido is not installed on the user's machine,
    fall back to a lightweight Matplotlib renderer so Excel/Word/PDF exports still contain charts.
    """
    if fig is None:
        return None
    try:
        if not getattr(fig, "data", None):
            return None
    except Exception:
        return None
    try:
        return fig.to_image(format="png", width=width, height=height, scale=scale)
    except Exception:
        return _fig_to_png_bytes_matplotlib(fig, width=width, height=height)


def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    try:
        return list(value)
    except Exception:
        return [value]


def _trace_name(trace: Any, fallback: str = "") -> str:
    name = getattr(trace, "name", None)
    return str(name or fallback or "Series")


def _fig_title(fig: go.Figure, default: str = "Biểu đồ") -> str:
    try:
        title = fig.layout.title.text
        return _clean_text(title) if title else default
    except Exception:
        return default


def _fig_y_title(fig: go.Figure) -> str:
    try:
        return str(fig.layout.yaxis.title.text or "")
    except Exception:
        return ""


def _fig_to_png_bytes_matplotlib(fig: go.Figure, width: int = 1200, height: int = 720) -> bytes | None:
    try:
        import matplotlib
        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt
        import numpy as np
    except Exception:
        return None
    try:
        px = 110
        figsize = (max(width / px, 6), max(height / px, 4))
        traces = list(getattr(fig, "data", []) or [])
        has_polar = any(str(getattr(t, "type", "")).lower() == "scatterpolar" for t in traces)
        if has_polar:
            fig_mpl = plt.figure(figsize=figsize)
            ax = fig_mpl.add_subplot(111, projection="polar")
        else:
            fig_mpl, ax = plt.subplots(figsize=figsize)
        fig_mpl.patch.set_facecolor("white")
        ax.set_title(_fig_title(fig), loc="left", fontsize=13, fontweight="bold", color="#064E47", pad=18)

        plotted = False
        if has_polar:
            for trace in traces:
                if str(getattr(trace, "type", "")).lower() != "scatterpolar":
                    continue
                r = [float(x) if x is not None and str(x) != "nan" else 0 for x in _as_list(getattr(trace, "r", []))]
                theta_raw = [str(x) for x in _as_list(getattr(trace, "theta", []))]
                # Drop duplicated closing point if present.
                if len(theta_raw) > 1 and theta_raw[0] == theta_raw[-1]:
                    theta_raw = theta_raw[:-1]
                    r = r[:-1]
                n = len(theta_raw)
                if not n:
                    continue
                angles = np.linspace(0, 2 * np.pi, n, endpoint=False).tolist()
                vals = r[:n]
                angles_closed = angles + angles[:1]
                vals_closed = vals + vals[:1]
                ax.plot(angles_closed, vals_closed, linewidth=2.2, marker="o", label=_trace_name(trace))
                ax.fill(angles_closed, vals_closed, alpha=0.16)
                ax.set_xticks(angles)
                ax.set_xticklabels([textwrap.shorten(t, width=18, placeholder="…") for t in theta_raw], fontsize=8)
                ax.set_ylim(0, 100)
                plotted = True
            ax.grid(True, alpha=0.35)
        else:
            for trace in traces:
                ttype = str(getattr(trace, "type", "")).lower()
                name = _trace_name(trace)
                if ttype == "bar":
                    orientation = str(getattr(trace, "orientation", "") or "").lower()
                    if orientation == "h":
                        x = [float(v) if v is not None and str(v) != "nan" else 0 for v in _as_list(getattr(trace, "x", []))]
                        y = [str(v) for v in _as_list(getattr(trace, "y", []))]
                        if x and y:
                            ax.barh(y, x, label=name)
                            plotted = True
                    else:
                        x = [str(v) for v in _as_list(getattr(trace, "x", []))]
                        y = [float(v) if v is not None and str(v) != "nan" else 0 for v in _as_list(getattr(trace, "y", []))]
                        if x and y:
                            ax.bar(x, y, label=name)
                            plotted = True
                elif ttype in {"scatter", "scattergl"}:
                    x = [str(v) for v in _as_list(getattr(trace, "x", []))]
                    y = [float(v) if v is not None and str(v) != "nan" else math.nan for v in _as_list(getattr(trace, "y", []))]
                    if x and y:
                        mode = str(getattr(trace, "mode", "lines") or "lines")
                        marker = "o" if "markers" in mode else None
                        linestyle = "-" if "lines" in mode or "markers" not in mode else "None"
                        ax.plot(x, y, linestyle=linestyle, marker=marker, linewidth=2.0, label=name)
                        plotted = True
                else:
                    x = [str(v) for v in _as_list(getattr(trace, "x", []))]
                    y = [float(v) if v is not None and str(v) != "nan" else 0 for v in _as_list(getattr(trace, "y", []))]
                    if x and y:
                        ax.plot(x, y, marker="o", linewidth=2.0, label=name)
                        plotted = True
            ax.grid(axis="y", alpha=0.25)
            ax.set_xlabel("Kỳ dữ liệu")
            ytitle = _fig_y_title(fig)
            if ytitle:
                ax.set_ylabel(ytitle)
            try:
                ax.tick_params(axis="x", rotation=30, labelsize=8)
                ax.tick_params(axis="y", labelsize=8)
            except Exception:
                pass
        if not plotted:
            plt.close(fig_mpl)
            return None
        if len(traces) > 1:
            ax.legend(loc="best", fontsize=8)
        fig_mpl.tight_layout()
        buf = BytesIO()
        fig_mpl.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig_mpl)
        return buf.getvalue()
    except Exception:
        try:
            plt.close("all")
        except Exception:
            pass
        return None


def _make_moat_radar_fig(moat_df: pd.DataFrame, ticker: str = "") -> go.Figure:
    df = _df(moat_df)
    if df.empty or "Nhóm Porter/Moat" not in df.columns:
        return go.Figure()
    if "Tỷ lệ đạt %" in df.columns:
        values = pd.to_numeric(df["Tỷ lệ đạt %"], errors="coerce").fillna(0).clip(0, 100).tolist()
    elif {"Điểm đạt", "Trọng số %"}.issubset(df.columns):
        values = (pd.to_numeric(df["Điểm đạt"], errors="coerce") / pd.to_numeric(df["Trọng số %"], errors="coerce").replace(0, pd.NA) * 100).fillna(0).clip(0, 100).tolist()
    else:
        values = []
    theta = df["Nhóm Porter/Moat"].astype(str).tolist()
    if not values:
        return go.Figure()
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=values + values[:1],
        theta=theta + theta[:1],
        fill="toself",
        name=ticker or "Moat score",
        hovertemplate="%{theta}<br>Điểm: %{r:.1f}/100<extra></extra>",
    ))
    fig.update_layout(
        title="Radar Porter Moat Score",
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        showlegend=False,
        height=520,
        margin=dict(l=50, r=50, t=70, b=45),
    )
    return fig


def _make_value_chain_radar_fig(value_chain_df: pd.DataFrame, ticker: str = "") -> go.Figure:
    df = _df(value_chain_df)
    if df.empty or "Hoạt động chuỗi giá trị" not in df.columns or "Điểm nhiệt" not in df.columns:
        return go.Figure()
    scores = pd.to_numeric(df["Điểm nhiệt"], errors="coerce").fillna(0).clip(0, 100).tolist()
    theta = df["Hoạt động chuỗi giá trị"].astype(str).tolist()
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=scores + scores[:1],
        theta=theta + theta[:1],
        fill="toself",
        name=ticker or "Chuỗi giá trị",
        hovertemplate="%{theta}<br>Điểm nhiệt: %{r:.1f}/100<extra></extra>",
    ))
    fig.update_layout(
        title="Radar chuỗi giá trị Porter",
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        showlegend=False,
        height=520,
        margin=dict(l=50, r=50, t=70, b=45),
    )
    return fig


def _make_peer_score_fig(peer_df: pd.DataFrame) -> go.Figure:
    df = _df(peer_df)
    if df.empty or "Mã" not in df.columns or "Điểm tổng hợp" not in df.columns:
        return go.Figure()
    df = df.copy()
    df["Điểm tổng hợp"] = pd.to_numeric(df["Điểm tổng hợp"], errors="coerce")
    df = df.dropna(subset=["Điểm tổng hợp"]).sort_values("Điểm tổng hợp", ascending=True).tail(12)
    if df.empty:
        return go.Figure()
    fig = go.Figure(go.Bar(
        x=df["Điểm tổng hợp"],
        y=df["Mã"].astype(str),
        orientation="h",
        hovertemplate="%{y}<br>Điểm tổng hợp: %{x:.1f}/100<extra></extra>",
    ))
    fig.update_layout(
        title="Xếp hạng so sánh doanh nghiệp cùng ngành",
        xaxis_title="Điểm tổng hợp /100",
        yaxis_title="Mã cổ phiếu",
        height=520,
        margin=dict(l=70, r=30, t=70, b=50),
        plot_bgcolor="rgba(0,0,0,0)",
        paper_bgcolor="rgba(0,0,0,0)",
    )
    return fig


def _build_chart_bundle(annual_df: pd.DataFrame, quarterly_df: pd.DataFrame, moat_df: pd.DataFrame,
                        value_chain_df: pd.DataFrame, peer_df: pd.DataFrame, ticker: str) -> list[tuple[str, go.Figure]]:
    charts: list[tuple[str, go.Figure]] = []
    annual = _df(annual_df)
    quarter = _df(quarterly_df)
    if not annual.empty:
        charts.extend([
            ("Tổng quan doanh nghiệp - Doanh thu và LNST năm + TTM", make_line_fig(annual, ["revenue_bil", "net_profit_bil"], "Doanh thu và lợi nhuận năm + TTM", "tỷ đồng")),
            ("Tổng quan doanh nghiệp - CFO, FCF và Owner Earnings năm + TTM", make_line_fig(annual, ["cfo_bil", "free_cash_flow_bil", "owner_earnings_bil"], "CFO, Free Cash Flow, Owner Earnings năm + TTM", "tỷ đồng")),
            ("Tổng quan doanh nghiệp - ROIC Operating Profit vs WACC năm + TTM", make_line_fig(annual, ["roic_operating_profit_pct", "wacc_pct"], "ROIC Operating Profit vs WACC năm + TTM", "%")),
            ("Tổng quan doanh nghiệp - EPS và OEPS năm + TTM", make_line_fig(annual, ["eps_vnd", "oeps_vnd"], "EPS và OEPS năm + TTM", "đồng/cp")),
            ("Tổng quan doanh nghiệp - FCF generation năm", make_fcf_generation_fig(annual, "FCF theo năm: LNTT -> FCF")),
            ("Tổng quan doanh nghiệp - Sử dụng dòng tiền năm", make_fcf_usage_fig(annual, "Sử dụng dòng tiền theo năm")),
            ("Tổng quan doanh nghiệp - FCF conversion năm", make_fcf_conversion_fig(annual, "FCF Conversion theo năm")),
            ("Tổng quan doanh nghiệp - DuPont profitability", make_dupont_profitability_fig(annual)),
            ("Tổng quan doanh nghiệp - DuPont driver", make_dupont_driver_fig(annual)),
            ("Tổng quan doanh nghiệp - ROIC và đầu tư", make_roic_investment_fig(annual)),
        ])
    if not quarter.empty:
        charts.extend([
            ("Tổng quan doanh nghiệp - Doanh thu và LNST quý", make_line_fig(quarter, ["revenue_bil", "net_profit_bil"], "Doanh thu và lợi nhuận theo quý", "tỷ đồng")),
            ("Tổng quan doanh nghiệp - CFO, FCF và Owner Earnings quý", make_line_fig(quarter, ["cfo_bil", "free_cash_flow_bil", "owner_earnings_bil"], "CFO, Free Cash Flow, Owner Earnings theo quý", "tỷ đồng")),
            ("Tổng quan doanh nghiệp - ROIC Operating Profit vs WACC quý", make_line_fig(quarter, ["roic_operating_profit_pct", "wacc_pct"], "ROIC Operating Profit vs WACC theo quý", "%")),
            ("Tổng quan doanh nghiệp - EPS và OEPS quý", make_line_fig(quarter, ["eps_vnd", "oeps_vnd"], "EPS và OEPS theo quý", "đồng/cp")),
        ])
    if not _df(moat_df).empty:
        charts.append(("Định giá chuyên sâu - Radar Porter Moat Score", _make_moat_radar_fig(moat_df, ticker)))
    if not _df(value_chain_df).empty:
        charts.append(("Định giá chuyên sâu - Radar chuỗi giá trị Porter", _make_value_chain_radar_fig(value_chain_df, ticker)))
    if not _df(peer_df).empty:
        charts.append(("So sánh doanh nghiệp - Điểm tổng hợp peer", _make_peer_score_fig(peer_df)))
    # Remove empty figures generated because a required column is absent.
    return [(title, fig) for title, fig in charts if getattr(fig, "data", None)]



def _build_module2_value_range_table(value_range: Any, target_mos_pct: float) -> pd.DataFrame:
    """Dải giá trị nội tại theo đúng logic tab Định giá chuyên sâu."""
    if value_range is None:
        return pd.DataFrame()
    weighted = getattr(value_range, "weighted_vnd", None)
    return pd.DataFrame([
        {"Chỉ tiêu": "Low", "Giá trị/cp": getattr(value_range, "low_vnd", None)},
        {"Chỉ tiêu": "Base/Median", "Giá trị/cp": getattr(value_range, "base_vnd", None)},
        {"Chỉ tiêu": "High", "Giá trị/cp": getattr(value_range, "high_vnd", None)},
        {"Chỉ tiêu": "Weighted", "Giá trị/cp": weighted},
        {"Chỉ tiêu": f"Giá mua theo MOS chọn {float(target_mos_pct):.0f}%", "Giá trị/cp": weighted * (1 - float(target_mos_pct) / 100) if weighted else None},
        {"Chỉ tiêu": "MOS hiện tại %", "Giá trị/cp": getattr(value_range, "mos_to_weighted_pct", None)},
        {"Chỉ tiêu": "MOS yêu cầu %", "Giá trị/cp": float(target_mos_pct)},
    ])


def _build_module2_strategic_assessment_table_export(company: CompanyOverview, annual: pd.DataFrame, cls: Any, value_range: Any, moat_df: pd.DataFrame, target_mos_pct: float) -> pd.DataFrame:
    """Bảng Đánh giá trọng yếu theo dữ liệu doanh nghiệp dùng cho báo cáo tổng hợp.

    Format lại riêng cột ``Số liệu/chứng cứ chính`` thay vì nhét số thô
    từ dataframe. Quy tắc: tỷ đồng và giá cổ phiếu 0 số thập phân, % và hệ số 1 số
    thập phân; thiếu dữ liệu hiển thị N/A. Không thay đổi phép tính định giá.
    """
    rows: list[dict[str, Any]] = []
    latest = annual.iloc[-1].to_dict() if annual is not None and not annual.empty else {}
    current_price = getattr(company, "current_price", None)
    weighted = getattr(value_range, "weighted_vnd", None) if value_range is not None else None
    mos_now = getattr(value_range, "mos_to_weighted_pct", None) if value_range is not None else None
    low = getattr(value_range, "low_vnd", None) if value_range is not None else None
    base = getattr(value_range, "base_vnd", None) if value_range is not None else None
    high = getattr(value_range, "high_vnd", None) if value_range is not None else None
    moat_score = moat_df.attrs.get("total_score", None) if isinstance(moat_df, pd.DataFrame) else None
    moat_level = moat_df.attrs.get("level", "N/A") if isinstance(moat_df, pd.DataFrame) else "N/A"

    roic = _latest_any(latest, ["roic_operating_profit_pct", "roic_standard_pct", "roic_pct"])
    roe = _latest_any(latest, ["roe_actual_pct", "roe_pct"])
    cfo_to_np = _latest_any(latest, ["cfo_to_net_profit", "cfo_to_np", "cfo_to_np_pct"])
    if cfo_to_np is None:
        cfo_to_np = _ratio_from_latest(latest, ["cfo_bil", "operating_cash_flow_bil"], ["net_profit_bil", "profit_after_tax_bil"])
    fcf_to_np = _latest_any(latest, ["fcf_to_net_profit", "fcf_to_np", "fcf_to_np_pct"])
    if fcf_to_np is None:
        fcf_to_np = _ratio_from_latest(latest, ["free_cash_flow_bil", "fcf_bil"], ["net_profit_bil", "profit_after_tax_bil"])

    buy_price = None
    if _num_or_none(weighted) is not None:
        buy_price = float(_num_or_none(weighted) or 0) * (1 - float(target_mos_pct) / 100)

    rows.append({
        "Nội dung cần đánh giá": "1. Doanh nghiệp thuộc loại nào?",
        "Kết luận theo mã": getattr(cls, "company_type", "N/A") if cls is not None else "N/A",
        "Số liệu/chứng cứ chính": (
            f"Độ tin cậy phân loại: {_fmt_score1(getattr(cls, 'confidence', None))}; "
            f"ngành: {_fmt_text_na(getattr(company, 'industry', None))}; "
            f"ROIC kỳ mới nhất: {_fmt_pct1(roic)}; ROE kỳ mới nhất: {_fmt_pct1(roe)}; "
            f"CFO/LNST: {_fmt_ratio1(cfo_to_np)}; FCF/LNST: {_fmt_ratio1(fcf_to_np)}."
        ),
        "Nguyên tắc áp dụng riêng": "Phân loại doanh nghiệp trước, chọn phương pháp định giá sau; không áp một công thức duy nhất cho mọi cổ phiếu.",
    })
    rows.append({
        "Nội dung cần đánh giá": "2. Lợi nhuận hiện tại có bền vững không?",
        "Kết luận theo mã": "Cần kiểm tra chất lượng lợi nhuận và dòng tiền",
        "Số liệu/chứng cứ chính": (
            f"Doanh thu: {_fmt_bil0(latest.get('revenue_bil'))}; "
            f"LNST: {_fmt_bil0(latest.get('net_profit_bil'))}; "
            f"CFO: {_fmt_bil0(latest.get('cfo_bil'))}; "
            f"FCF: {_fmt_bil0(latest.get('free_cash_flow_bil'))}; "
            f"Owner Earnings: {_fmt_bil0(latest.get('owner_earnings_bil'))}; "
            f"CFO/LNST: {_fmt_ratio1(cfo_to_np)}."
        ),
        "Nguyên tắc áp dụng riêng": "Lợi nhuận phải đi kèm dòng tiền thật; nếu CFO/FCF yếu kéo dài, cần haircut định giá và tăng MOS.",
    })
    rows.append({
        "Nội dung cần đánh giá": "3. Moat đến từ đâu?",
        "Kết luận theo mã": f"{moat_level} ({_fmt_score1(moat_score)})",
        "Số liệu/chứng cứ chính": (
            f"Biên gộp: {_fmt_pct1(latest.get('gross_margin_pct'))}; "
            f"biên EBIT: {_fmt_pct1(latest.get('ebit_margin_pct'))}; "
            f"biên ròng: {_fmt_pct1(latest.get('net_margin_pct'))}; "
            f"vòng quay tài sản: {_fmt_ratio1(latest.get('asset_turnover'))}; "
            f"ROIC Operating Profit: {_fmt_pct1(roic)}."
        ),
        "Nguyên tắc áp dụng riêng": "Porter moat phải bám hoạt động tạo giá trị: chi phí thấp, khác biệt hóa, phân phối, thương hiệu, switching cost hoặc tài sản đặc thù.",
    })
    rows.append({
        "Nội dung cần đánh giá": "4. ROIC/ROCE cao là moat thật hay chu kỳ?",
        "Kết luận theo mã": "Cần đối chiếu ROIC với WACC, CFO/FCF và chu kỳ ngành",
        "Số liệu/chứng cứ chính": (
            f"ROIC: {_fmt_pct1(roic)}; "
            f"WACC: {_fmt_pct1(latest.get('wacc_pct'))}; "
            f"Deployed Capital: {_fmt_bil0(latest.get('deployed_capital_bil'))}; "
            f"Tổng tài sản: {_fmt_bil0(latest.get('total_assets_bil'))}; "
            f"Vốn chủ: {_fmt_bil0(latest.get('equity_bil'))}."
        ),
        "Nguyên tắc áp dụng riêng": "ROIC cao chỉ đáng giá nếu bền vững và tái đầu tư được; ROIC chu kỳ phải chuẩn hóa qua nhiều năm.",
    })
    rows.append({
        "Nội dung cần đánh giá": "5. Giá hiện tại có đủ biên an toàn không?",
        "Kết luận theo mã": getattr(value_range, "recommendation", "N/A") if value_range is not None else "N/A",
        "Số liệu/chứng cứ chính": (
            f"Giá hiện tại: {_fmt_vnd0(current_price)}; "
            f"giá trị thấp: {_fmt_vnd0(low)}; "
            f"giá trị cơ sở: {_fmt_vnd0(base)}; "
            f"giá trị cao: {_fmt_vnd0(high)}; "
            f"giá trị weighted: {_fmt_vnd0(weighted)}; "
            f"MOS hiện tại: {_fmt_pct1(mos_now)}; "
            f"MOS yêu cầu: {float(target_mos_pct):,.1f}%; "
            f"giá mua tối đa theo MOS: {_fmt_vnd0(buy_price)}."
        ),
        "Nguyên tắc áp dụng riêng": "Graham/Li Lu/Howard Marks: mua dưới giá trị nội tại với biên an toàn đủ lớn; khi dữ liệu hoặc moat chưa rõ phải tăng MOS.",
    })
    return pd.DataFrame(rows)


def _df_columns_existing(df: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    src = _df(df)
    cols = [c for c in columns if c in src.columns]
    return src[cols].copy() if cols else pd.DataFrame()


def _make_scorecard_radar_fig(scorecard: pd.DataFrame, title: str, label_col: str = "Nhóm tiêu chí") -> go.Figure:
    """Radar chart for scorecard tables so the consolidated report keeps the app's scoring visuals."""
    df = _df(scorecard)
    if df.empty:
        return go.Figure()
    if label_col not in df.columns:
        # Support Định giá chuyên sâu Porter scorecards as well.
        if "Nhóm Porter/Moat" in df.columns:
            label_col = "Nhóm Porter/Moat"
        elif "Hoạt động chuỗi giá trị" in df.columns:
            label_col = "Hoạt động chuỗi giá trị"
        else:
            return go.Figure()
    chart_df = df.copy()
    label_series = chart_df[label_col].astype(str)
    # Remove total rows from radar because they distort the shape.
    mask_total = label_series.str.contains("TỔNG|Tổng", case=False, na=False)
    chart_df = chart_df.loc[~mask_total].copy()
    if chart_df.empty:
        return go.Figure()
    if "Tỷ lệ đạt" in chart_df.columns:
        chart_df["_heat"] = pd.to_numeric(chart_df["Tỷ lệ đạt"], errors="coerce")
    elif {"Điểm", "Trọng số"}.issubset(chart_df.columns):
        score = pd.to_numeric(chart_df["Điểm"], errors="coerce")
        weight = pd.to_numeric(chart_df["Trọng số"], errors="coerce").replace(0, pd.NA)
        chart_df["_heat"] = score / weight * 100
    elif {"Điểm đạt", "Trọng số %"}.issubset(chart_df.columns):
        score = pd.to_numeric(chart_df["Điểm đạt"], errors="coerce")
        weight = pd.to_numeric(chart_df["Trọng số %"], errors="coerce").replace(0, pd.NA)
        chart_df["_heat"] = score / weight * 100
    elif "Điểm nhiệt" in chart_df.columns:
        chart_df["_heat"] = pd.to_numeric(chart_df["Điểm nhiệt"], errors="coerce")
    else:
        return go.Figure()
    chart_df["_heat"] = chart_df["_heat"].fillna(0).clip(0, 100)
    chart_df[label_col] = chart_df[label_col].astype(str).str.replace("Tỷ lệ ", "", regex=False).str.slice(0, 38)
    theta = chart_df[label_col].tolist()
    r = chart_df["_heat"].astype(float).tolist()
    if len(theta) < 3:
        return go.Figure()
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(
        r=r + r[:1],
        theta=theta + theta[:1],
        fill="toself",
        name="Điểm nhiệt",
        hovertemplate="%{theta}<br>Điểm nhiệt: %{r:.1f}/100<extra></extra>",
    ))
    fig.update_layout(
        title=title,
        polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
        showlegend=False,
        height=520,
        margin=dict(l=50, r=50, t=70, b=45),
    )
    return fig


def _scorecard_total_table(scorecard: pd.DataFrame, label: str) -> pd.DataFrame:
    """Compact non-empty summary for every scoring/evaluation table."""
    df = _df(scorecard)
    if df.empty:
        return pd.DataFrame([{"Bảng đánh giá/chấm điểm": label, "Tình trạng": "Chưa có dữ liệu", "Tổng điểm/đánh giá": "N/A", "Điểm mạnh/yếu cần xem": "Cần cập nhật dữ liệu"}])
    total_text = "N/A"
    if "Nhóm tiêu chí" in df.columns and "Điểm" in df.columns:
        total = df[df["Nhóm tiêu chí"].astype(str).str.contains("TỔNG", case=False, na=False)]
        if not total.empty:
            r0 = total.iloc[0]
            total_text = f"{r0.get('Điểm', 'N/A')}/{r0.get('Trọng số', 'N/A')} - {r0.get('Tín hiệu', '')}"
        else:
            total_text = f"{pd.to_numeric(df.get('Điểm'), errors='coerce').sum():.1f}/{pd.to_numeric(df.get('Trọng số'), errors='coerce').sum():.1f}"
    elif "Nhóm Porter/Moat" in df.columns and "Điểm đạt" in df.columns:
        total_text = f"{df.attrs.get('total_score', pd.to_numeric(df.get('Điểm đạt'), errors='coerce').sum()):.1f}/100 - {df.attrs.get('level', 'N/A')}"
    elif "Điểm nhiệt" in df.columns:
        avg = pd.to_numeric(df.get("Điểm nhiệt"), errors="coerce").mean()
        total_text = f"Điểm nhiệt TB {avg:.1f}/100" if pd.notna(avg) else "N/A"
    note_cols = [c for c in ["Nhận xét tự động", "Diễn giải", "Tín hiệu", "Đánh giá sơ bộ", "Mức độ"] if c in df.columns]
    notes = []
    for _, row in df.head(6).iterrows():
        part = []
        for c in note_cols[:2]:
            v = str(row.get(c, "")).strip()
            if v and v.lower() != "nan":
                part.append(v)
        if part:
            notes.append("; ".join(part))
    return pd.DataFrame([{
        "Bảng đánh giá/chấm điểm": label,
        "Tình trạng": "Có dữ liệu",
        "Tổng điểm/đánh giá": total_text,
        "Điểm mạnh/yếu cần xem": " | ".join(notes[:3]) if notes else "Xem chi tiết trong bảng bên dưới",
    }])


def _build_scorecard_inventory_table(items: list[tuple[str, pd.DataFrame]]) -> pd.DataFrame:
    frames = [_scorecard_total_table(df, name) for name, df in items]
    return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()


def _build_value_chain_summary_table(value_chain_df: pd.DataFrame) -> pd.DataFrame:
    df = _df(value_chain_df)
    if df.empty:
        return pd.DataFrame()
    score = pd.to_numeric(df.get("Điểm nhiệt"), errors="coerce") if "Điểm nhiệt" in df.columns else pd.Series(dtype=float)
    avg = score.mean() if not score.empty else pd.NA
    strong = []
    weak = []
    if "Đánh giá sơ bộ" in df.columns and "Hoạt động chuỗi giá trị" in df.columns:
        strong = df[df["Đánh giá sơ bộ"].astype(str).str.contains("Mạnh|Thuận lợi", case=False, na=False)]["Hoạt động chuỗi giá trị"].astype(str).head(5).tolist()
        weak = df[df["Đánh giá sơ bộ"].astype(str).str.contains("Yếu|Cần|Bất lợi|Chưa", case=False, na=False)]["Hoạt động chuỗi giá trị"].astype(str).head(5).tolist()
    if pd.isna(avg):
        level = "Theo dõi"
    elif avg >= 70:
        level = "Mạnh"
    elif avg >= 55:
        level = "Trung bình/khá"
    else:
        level = "Cần kiểm tra"
    return pd.DataFrame([{
        "Nội dung": "Đánh giá tổng hợp chuỗi giá trị Porter",
        "Điểm nhiệt TB": round(float(avg), 1) if pd.notna(avg) else "N/A",
        "Mức đánh giá": level,
        "Điểm mạnh nổi bật": ", ".join(strong) if strong else "Chưa xác định rõ",
        "Điểm yếu/cần kiểm tra": ", ".join(weak) if weak else "Chưa xác định rõ",
    }])


def _hide_financial_manipulation_logic_cols(df: pd.DataFrame, *, drop_layer: bool = False) -> pd.DataFrame:
    """Hide source/logic columns from financial manipulation tables in the consolidated report.

    drop_layer=True is used for the individual layer detail tables because their section title
    already states the layer name; keeping the column wastes width and can squeeze the Kỳ column.
    """
    src = _df(df).copy()
    if src.empty:
        return src
    hidden = {"Nguồn/logic", "Nguồn / logic", "source", "Source", "Nguồn", "Nguồn dữ liệu"}
    if drop_layer:
        hidden.add("Lớp")
    return src.drop(columns=[c for c in src.columns if str(c) in hidden], errors="ignore")


def _latest_layer_summary(layer_name: str, df: pd.DataFrame, metric_candidates: list[str]) -> dict[str, Any]:
    src = _hide_financial_manipulation_logic_cols(df)
    if src.empty:
        return {
            "Lớp": layer_name,
            "Kỳ mới nhất": "N/A",
            "Chỉ tiêu chính": "N/A",
            "Giá trị": "N/A",
            "Mức cảnh báo": "Chưa đủ dữ liệu",
            "Tín hiệu": "Chưa đủ dữ liệu để tính lớp cảnh báo này.",
            "Cần kiểm tra": "Bổ sung dữ liệu BCTC theo năm, đặc biệt doanh thu, phải thu, tài sản, CFO, hàng tồn kho, chi phí và khấu hao.",
        }
    latest = src.iloc[-1].to_dict()
    metric_name = next((m for m in metric_candidates if m in latest), metric_candidates[0] if metric_candidates else "Chỉ tiêu")
    return {
        "Lớp": layer_name,
        "Kỳ mới nhất": latest.get("Kỳ", latest.get("period", "N/A")),
        "Chỉ tiêu chính": metric_name,
        "Giá trị": latest.get(metric_name, "N/A"),
        "Mức cảnh báo": latest.get("Mức cảnh báo", "N/A"),
        "Tín hiệu": latest.get("Tín hiệu", latest.get("Nhận xét", "N/A")),
        "Cần kiểm tra": latest.get("Cần kiểm tra", latest.get("Biến nổi bật/cần kiểm tra", latest.get("Biến thiếu/cần kiểm tra", "N/A"))),
    }


def _build_financial_manipulation_summary_table(
    beneish_df: pd.DataFrame,
    accrual_quality_df: pd.DataFrame,
    modified_jones_df: pd.DataFrame,
    rem_df: pd.DataFrame,
) -> pd.DataFrame:
    """Summary table for the consolidated report: one latest-period row for each manipulation layer."""
    rows = [
        _latest_layer_summary("1. Beneish M-Score", beneish_df, ["M-Score"]),
        _latest_layer_summary("2. Accrual Quality/Sloan", accrual_quality_df, ["Sloan accrual ratio", "CFO/LNST", "FCF/LNST"]),
        _latest_layer_summary("3. Modified Jones/Kothari", modified_jones_df, ["DA Modified Jones", "DA Kothari"]),
        _latest_layer_summary("4. REM - hoạt động thật", rem_df, ["REM Score", "Abnormal CFO", "Abnormal PROD", "Abnormal DISEXP"]),
    ]
    return pd.DataFrame(rows)


def build_report_package(
    company: CompanyOverview,
    annual_df: pd.DataFrame,
    quarterly_df: pd.DataFrame,
    *,
    valuation_df: pd.DataFrame | None = None,
    moat_df: pd.DataFrame | None = None,
    value_chain_df: pd.DataFrame | None = None,
    scenario_df: pd.DataFrame | None = None,
    peer_df: pd.DataFrame | None = None,
    web_df: pd.DataFrame | None = None,
    assumptions: dict[str, Any] | None = None,
    source_label: str = "",
    paths: Iterable[Any] | None = None,
    target_mos_pct: float = 50.0,
    module1_summary: str | None = None,
    module1_value_investing_summary: str | None = None,
    module1_mos_summary: str | None = None,
    module2_summary: str | None = None,
) -> ReportPackage:
    warnings: list[str] = []
    ticker = _safe_ticker(getattr(company, "ticker", ""))
    annual = _df(annual_df)
    quarterly = _df(quarterly_df)
    peer = _df(peer_df)
    web = _df(web_df)

    if assumptions is None:
        try:
            assumptions = load_assumptions(DEFAULT_ASSUMPTIONS_PATH)
        except Exception as exc:
            warnings.append(f"Không tải được assumptions Định giá chuyên sâu: {exc}")
            assumptions = {}
    assumptions = dict(assumptions or {})
    if "target_mos_pct" not in assumptions:
        assumptions["target_mos_pct"] = float(target_mos_pct)

    module1_mos_df = build_mos_valuation_table(company, annual, mos_rate=float(target_mos_pct) / 100) if not annual.empty else pd.DataFrame()
    ratio_scorecard = build_financial_ratio_scorecard(annual) if not annual.empty else pd.DataFrame()
    cashflow_scorecard = build_cashflow_scorecard(annual) if not annual.empty else pd.DataFrame()
    combined_assessment = build_combined_assessment_table(company, annual, quarterly, module1_mos_df) if not annual.empty else pd.DataFrame()

    if valuation_df is None:
        valuation_df = build_module2_valuation_table(company, annual, assumptions) if not annual.empty else pd.DataFrame()
    if moat_df is None:
        moat_df = build_porter_moat_scorecard(company, annual)
    if value_chain_df is None:
        value_chain_df = build_value_chain_table(company, annual)
    current_price = getattr(company, "current_price", None)
    value_range = build_valuation_range(_df(valuation_df), current_price, float(target_mos_pct)) if not _df(valuation_df).empty else None
    if scenario_df is None:
        scenario_df = build_risk_scenario_table(company, annual, value_range) if value_range is not None and not annual.empty else pd.DataFrame()
    cls = classify_company(company, annual) if not annual.empty else None

    if module1_summary is None:
        module1_summary = build_quick_summary(company, annual_df=annual) if not annual.empty else "Chưa đủ dữ liệu để tạo tóm tắt Tổng quan doanh nghiệp."
    if module1_value_investing_summary is None:
        module1_value_investing_summary = build_value_investing_assessment(company, annual, ratio_scorecard) if not annual.empty else "Chưa đủ dữ liệu để nhận xét theo triết lý đầu tư giá trị."
    if module1_mos_summary is None:
        module1_mos_summary = build_mos_detailed_summary(module1_mos_df) if not module1_mos_df.empty else "Chưa đủ dữ liệu để định giá MOS Tổng quan doanh nghiệp."
    if module2_summary is None:
        module2_summary = build_module2_summary(company, annual, _df(valuation_df), _df(moat_df)) if not annual.empty else "Chưa đủ dữ liệu để tạo tóm tắt Định giá chuyên sâu."

    paths_text = "\n".join([str(p) for p in (paths or []) if p is not None]) or "N/A"
    metrics = build_metric_dict(company)
    latest_cards = latest_metric_cards(annual) if not annual.empty else {}
    flags = build_flags(company, annual_df=annual, quarterly_df=quarterly) if not annual.empty else []
    flags_df = pd.DataFrame(flags) if flags else pd.DataFrame()
    overview_df = pd.DataFrame([
        {"Chỉ tiêu": "Mã", "Giá trị": getattr(company, "ticker", "")},
        {"Chỉ tiêu": "Tên doanh nghiệp", "Giá trị": getattr(company, "company_name", "")},
        {"Chỉ tiêu": "Sàn", "Giá trị": getattr(company, "exchange", "")},
        {"Chỉ tiêu": "Ngành", "Giá trị": getattr(company, "industry", "")},
        {"Chỉ tiêu": "Phân ngành", "Giá trị": getattr(company, "sub_industry", "")},
        {"Chỉ tiêu": "Giá hiện tại", "Giá trị": getattr(company, "current_price", "")},
        {"Chỉ tiêu": "Cập nhật", "Giá trị": getattr(company, "updated_at", "")},
        {"Chỉ tiêu": "Chế độ dữ liệu", "Giá trị": "Dữ liệu nội bộ"},
    ])
    metrics_df = pd.DataFrame([{"Chỉ tiêu": k, "Giá trị": v} for k, v in {**metrics, **latest_cards}.items()])

    figures = _build_chart_bundle(annual, quarterly, _df(moat_df), _df(value_chain_df), peer, ticker)
    if not figures:
        warnings.append("Không có biểu đồ nào đủ dữ liệu để xuất.")

    # V23.40: báo cáo tổng hợp chuyển từ gom vài bảng chính sang render theo từng tab của từng phần.
    # Mục tiêu là khi in/Save as PDF, người dùng nhìn thấy đủ đánh giá, chấm điểm, phân tích,
    # biểu đồ và số liệu như các tab trong app, thay vì chỉ nhận một bản export rút gọn.
    value_range_df = _build_module2_value_range_table(value_range, float(target_mos_pct))
    strategic_assessment = _build_module2_strategic_assessment_table_export(company, annual, cls, value_range, _df(moat_df), float(target_mos_pct))
    classification_reasons = getattr(cls, "reasons", []) if cls is not None else []
    classification_reason_text = "\n".join([f"- {r}" for r in classification_reasons if str(r).strip()]) or "Chưa có đủ dữ liệu để giải thích phân loại."

    # Tổng quan doanh nghiệp extra tables matching dashboard tabs.
    fcf_scorecard_year = build_cashflow_scorecard(annual) if not annual.empty else pd.DataFrame()
    fcf_scorecard_quarter = build_cashflow_scorecard(quarterly) if not quarterly.empty else pd.DataFrame()
    fcf_alerts_year = build_cashflow_situation_alerts(annual) if not annual.empty else pd.DataFrame()
    fcf_alerts_quarter = build_cashflow_situation_alerts(quarterly) if not quarterly.empty else pd.DataFrame()
    fcf_table_year = build_fcf_analysis_table(annual) if not annual.empty else pd.DataFrame()
    fcf_table_quarter = build_fcf_analysis_table(quarterly) if not quarterly.empty else pd.DataFrame()
    ratio_table_year = build_financial_ratio_table(annual) if not annual.empty else pd.DataFrame()
    ratio_table_quarter = build_financial_ratio_table(quarterly) if not quarterly.empty else pd.DataFrame()
    ratio_alerts = build_financial_ratio_alerts(annual) if not annual.empty else pd.DataFrame()
    dupont_table = _df_columns_existing(annual if not annual.empty else quarterly, [
        "period", "roe_pct", "roa_pct", "gross_margin_pct", "net_margin_pct", "asset_turnover", "equity_multiplier", "roe_actual_pct", "roe_dupont_pct"
    ])
    roic_table = _df_columns_existing(annual if not annual.empty else quarterly, [
        "period", "core_operating_profit_bil", "nopat_bil", "deployed_capital_bil", "avg_deployed_capital_bil",
        "interest_bearing_debt_bil", "avg_interest_bearing_debt_bil", "market_cap_bil", "equity_weight_pct",
        "debt_weight_pct", "cost_of_equity_pct", "cost_of_debt_pct", "after_tax_cost_of_debt_pct", "tax_rate_pct",
        "beta", "roic_operating_profit_pct", "wacc_pct", "wacc_quality", "expansion_investment_bil",
        "inventory_change_bil", "investment_subsidiary_bil", "total_investment_bil"
    ])
    latest_cards_df = pd.DataFrame([{"Chỉ tiêu": k, "Giá trị": v} for k, v in latest_cards.items()])

    scorecard_items = [
        ("Tổng quan doanh nghiệp - FCF & dòng tiền theo năm", fcf_scorecard_year),
        ("Tổng quan doanh nghiệp - FCF & dòng tiền theo quý", fcf_scorecard_quarter),
        ("Tổng quan doanh nghiệp - Phân tích chỉ số tài chính 100 điểm", ratio_scorecard),
        ("Định giá chuyên sâu - Porter Moat Score", _df(moat_df)),
        ("Định giá chuyên sâu - Chuỗi giá trị Porter", _df(value_chain_df)),
    ]
    scorecard_inventory_df = _build_scorecard_inventory_table(scorecard_items)
    value_chain_summary_df = _build_value_chain_summary_table(_df(value_chain_df))

    # V23.57: add the Financial Manipulation tab into the consolidated full report.
    # All tables are formatted by _format_print_df later, so tỷ đồng/%, ratios follow the project display rules.
    beneish_report_df = build_beneish_mscore_table(company, annual) if not annual.empty else pd.DataFrame()
    accrual_quality_report_df = build_accrual_quality_table(company, annual) if not annual.empty else pd.DataFrame()
    modified_jones_report_df = build_modified_jones_kothari_table(company, annual) if not annual.empty else pd.DataFrame()
    rem_report_df = build_real_earnings_management_table(company, annual) if not annual.empty else pd.DataFrame()
    financial_manipulation_summary_df = _build_financial_manipulation_summary_table(
        beneish_report_df,
        accrual_quality_report_df,
        modified_jones_report_df,
        rem_report_df,
    )
    beneish_report_df = _hide_financial_manipulation_logic_cols(beneish_report_df, drop_layer=True)
    accrual_quality_report_df = _hide_financial_manipulation_logic_cols(accrual_quality_report_df, drop_layer=True)
    modified_jones_report_df = _hide_financial_manipulation_logic_cols(modified_jones_report_df, drop_layer=True)
    rem_report_df = _hide_financial_manipulation_logic_cols(rem_report_df, drop_layer=True)

    # V23.40: đưa các biểu đồ mạng nhện/chấm điểm vào báo cáo tổng hợp để không thiếu phần đánh giá.
    extra_score_figures = [
        ("Tổng quan doanh nghiệp - Radar chấm điểm FCF & dòng tiền theo năm", _make_scorecard_radar_fig(fcf_scorecard_year, "Radar chấm điểm FCF & dòng tiền theo năm")),
        ("Tổng quan doanh nghiệp - Radar chấm điểm FCF & dòng tiền theo quý", _make_scorecard_radar_fig(fcf_scorecard_quarter, "Radar chấm điểm FCF & dòng tiền theo quý")),
        ("Tổng quan doanh nghiệp - Radar chấm điểm chỉ số tài chính", _make_scorecard_radar_fig(ratio_scorecard, "Radar chấm điểm phân tích chỉ số tài chính")),
    ]
    figures.extend([(t, f) for t, f in extra_score_figures if getattr(f, "data", None)])

    title = f"Báo cáo tổng hợp toàn bộ nội dung - {ticker}"
    sections = [
        ReportSection(
            "0. Bìa báo cáo và thông tin doanh nghiệp",
            [
                f"Báo cáo được tạo tự động từ dashboard Trecapital cho mã {ticker}.",
                "Báo cáo này tổng hợp toàn bộ từng tab của từng phần: đánh giá, chấm điểm, phân tích, biểu đồ và bảng số liệu chính. Luồng xuất duy nhất là mở trang Báo cáo tổng hợp toàn bộ nội dung rồi in/Save as PDF để giữ format giống app.",
            ],
            [("Thông tin doanh nghiệp", overview_df), ("KPI tổng hợp", metrics_df), ("KPI kỳ gần nhất", latest_cards_df), ("Cảnh báo nhanh", flags_df)],
            [],
        ),
        ReportSection(
            "0.1 Tổng hợp các bảng đánh giá và chấm điểm trong báo cáo",
            [
                "Mục kiểm tra nhanh: toàn bộ bảng đánh giá, chấm điểm, cảnh báo và định giá quan trọng phải xuất hiện ở đây trước khi được trình bày lại trong từng tab tương ứng bên dưới."
            ],
            [
                ("Danh mục bảng đánh giá/chấm điểm", scorecard_inventory_df),
                ("Tổng quan doanh nghiệp - Cảnh báo / điểm cần kiểm tra", combined_assessment),
                ("Tổng quan doanh nghiệp - Kết quả định giá MOS", module1_mos_df),
                ("Tổng quan doanh nghiệp - Bộ tiêu chí đánh giá FCF & dòng tiền theo năm", fcf_scorecard_year),
                ("Tổng quan doanh nghiệp - Cảnh báo dòng tiền theo năm", fcf_alerts_year),
                ("Tổng quan doanh nghiệp - Bộ tiêu chí đánh giá FCF & dòng tiền theo quý", fcf_scorecard_quarter),
                ("Tổng quan doanh nghiệp - Cảnh báo dòng tiền theo quý", fcf_alerts_quarter),
                ("Tổng quan doanh nghiệp - Bộ tiêu chí đánh giá tự động - 100 điểm", ratio_scorecard),
                ("Tổng quan doanh nghiệp - Cảnh báo/tình huống chỉ số tài chính", ratio_alerts),
                ("Định giá chuyên sâu - Đánh giá trọng yếu theo dữ liệu doanh nghiệp", strategic_assessment),
                ("Định giá chuyên sâu - Dải giá trị nội tại", value_range_df),
                ("Định giá chuyên sâu - Bảng định giá theo từng phương pháp", _df(valuation_df)),
                ("Định giá chuyên sâu - Tổng hợp chấm điểm Porter Moat", _scorecard_total_table(_df(moat_df), "Định giá chuyên sâu - Porter Moat Score")),
                ("Định giá chuyên sâu - Bảng điểm lợi thế cạnh tranh theo Porter", _df(moat_df)),
                ("Định giá chuyên sâu - Đánh giá tổng hợp chuỗi giá trị Porter", value_chain_summary_df),
                ("Định giá chuyên sâu - Bản đồ chuỗi giá trị theo Porter", _df(value_chain_df)),
                ("Định giá chuyên sâu - Kịch bản & rủi ro", _df(scenario_df)),
                ("So sánh doanh nghiệp - Kết quả peer comparison", peer),
            ],
            [],
        ),

        # MODULE 1 - từng tab
        ReportSection(
            "1.1 Tổng quan doanh nghiệp / Tab Tóm tắt",
            [module1_summary, module1_value_investing_summary, module1_mos_summary],
            [("Kết quả định giá MOS", module1_mos_df), ("Cảnh báo / điểm cần kiểm tra", combined_assessment)],
            [],
        ),
        ReportSection(
            "1.2 Tổng quan doanh nghiệp / Tab Biểu đồ tài chính",
            ["Biểu đồ tài chính năm + TTM và 20 quý: doanh thu, lợi nhuận, CFO, FCF, Owner Earnings, ROIC/WACC, EPS/OEPS."],
            [],
            [(t, f) for t, f in figures if t.startswith("Tổng quan doanh nghiệp - Doanh thu") or t.startswith("Tổng quan doanh nghiệp - CFO") or t.startswith("Tổng quan doanh nghiệp - ROIC Operating") or t.startswith("Tổng quan doanh nghiệp - EPS")],
        ),
        ReportSection(
            "1.3 Tổng quan doanh nghiệp / Tab FCF & dòng tiền",
            ["Phân tích quá trình LNTT → điều chỉnh phi tiền mặt → thay đổi vốn lưu động → Capex → FCF, sau đó đánh giá cách doanh nghiệp dùng FCF cho trả nợ, cổ tức, đầu tư và tăng/giảm tiền."],
            [
                ("Bộ tiêu chí đánh giá FCF & dòng tiền theo năm", fcf_scorecard_year),
                ("Cảnh báo dòng tiền theo năm", fcf_alerts_year),
                ("Bảng phân tích sử dụng dòng tiền theo năm", fcf_table_year),
                ("Bộ tiêu chí đánh giá FCF & dòng tiền theo quý", fcf_scorecard_quarter),
                ("Cảnh báo dòng tiền theo quý", fcf_alerts_quarter),
                ("Bảng phân tích sử dụng dòng tiền theo quý", fcf_table_quarter),
            ],
            [(t, f) for t, f in figures if "FCF generation" in t or "Sử dụng dòng tiền" in t or "FCF conversion" in t or "Radar chấm điểm FCF" in t],
        ),
        ReportSection(
            "1.4 Tổng quan doanh nghiệp / Tab Phân tích chỉ số TC",
            [module1_value_investing_summary],
            [
                ("Bộ tiêu chí đánh giá tự động - 100 điểm", ratio_scorecard),
                ("Cảnh báo/tình huống chỉ số tài chính", ratio_alerts),
                ("Bảng phân tích chỉ số tài chính theo năm", ratio_table_year),
                ("Bảng phân tích chỉ số tài chính theo quý", ratio_table_quarter),
            ],
            [(t, f) for t, f in figures if "Radar chấm điểm chỉ số tài chính" in t],
        ),
        ReportSection(
            "1.5 Tổng quan doanh nghiệp / Tab DuPont",
            ["Phân rã ROE/ROA theo biên lợi nhuận, vòng quay tài sản và đòn bẩy để nhận diện động lực sinh lời."],
            [("Bảng chỉ tiêu DuPont", dupont_table)],
            [(t, f) for t, f in figures if "DuPont" in t],
        ),
        ReportSection(
            "1.6 Tổng quan doanh nghiệp / Tab ROIC & đầu tư",
            ["Đối chiếu ROIC Operating Profit với WACC và kiểm tra lượng vốn tái đầu tư để xem doanh nghiệp có tạo giá trị trên vốn sử dụng hay không."],
            [("Bảng ROIC & đầu tư", roic_table)],
            [(t, f) for t, f in figures if "ROIC và đầu tư" in t],
        ),
        # MODULE 2 - từng tab
        ReportSection(
            "2.1 Định giá chuyên sâu / Tab Định giá chuyên sâu",
            [
                module2_summary,
                f"Phân loại sơ bộ: {getattr(cls, 'company_type', 'N/A') if cls else 'N/A'}; độ tin cậy: {getattr(cls, 'confidence', 0):.0f}/100." if cls else "Chưa phân loại được do thiếu dữ liệu.",
                f"MOS yêu cầu đang áp dụng: {float(target_mos_pct):.0f}%.",
                f"Khuyến nghị theo dải định giá: {getattr(value_range, 'recommendation', 'N/A') if value_range else 'N/A'}.",
            ],
            [
                ("Đánh giá trọng yếu theo dữ liệu doanh nghiệp", strategic_assessment),
                ("Dải giá trị nội tại", value_range_df),
                ("Bảng định giá theo từng phương pháp", _df(valuation_df)),
            ],
            [],
        ),
        ReportSection(
            "2.2 Định giá chuyên sâu / Tab Porter Moat Score",
            [
                f"Tổng điểm Porter Moat: {_df(moat_df).attrs.get('total_score', moat_df.attrs.get('total_score', 'N/A') if isinstance(moat_df, pd.DataFrame) else 'N/A') if isinstance(moat_df, pd.DataFrame) else 'N/A'} - {moat_df.attrs.get('level', 'N/A') if isinstance(moat_df, pd.DataFrame) else 'N/A'}.",
                "Lý do phân loại:\n" + classification_reason_text,
            ],
            [("Tổng hợp chấm điểm Porter Moat", _scorecard_total_table(_df(moat_df), "Định giá chuyên sâu - Porter Moat Score")), ("Bảng điểm lợi thế cạnh tranh theo Porter", _df(moat_df))],
            [(t, f) for t, f in figures if "Porter Moat" in t],
        ),
        ReportSection(
            "2.3 Định giá chuyên sâu / Tab Chuỗi giá trị",
            ["Bản đồ chuỗi giá trị theo Porter: mỗi hoạt động được liên kết với tín hiệu định lượng hiện có và bằng chứng định tính cần kiểm tra thêm."],
            [("Đánh giá tổng hợp chuỗi giá trị Porter", value_chain_summary_df), ("Bản đồ chuỗi giá trị theo Porter", _df(value_chain_df))],
            [(t, f) for t, f in figures if "chuỗi giá trị" in t.lower()],
        ),
        ReportSection(
            "2.4 Định giá chuyên sâu / Tab Kịch bản & rủi ro",
            ["Kịch bản rủi ro giúp kiểm tra downside/upside, điểm cần kiểm tra và tác động đến biên an toàn."],
            [("Kịch bản & rủi ro", _df(scenario_df))],
            [],
        ),
        ReportSection(
            "2.5 Định giá chuyên sâu / Tab Thao túng tài chính",
            [
                "Nguyên tắc: app không kết luận doanh nghiệp gian lận. Bốn lớp dưới đây chỉ tạo cờ đỏ định lượng để kiểm tra sâu chất lượng BCTC: lợi nhuận có đi kèm tiền thật không, accruals có bất thường không, doanh thu/phải thu/tài sản có bị làm đẹp không và có dấu hiệu quản trị lợi nhuận qua hoạt động thật không.",
                "Cách dùng trong app: khi một hoặc nhiều lớp cảnh báo cao, hãy giảm độ tin cậy của lợi nhuận kế toán khi định giá; đọc kỹ thuyết minh doanh thu, phải thu, tồn kho, khấu hao, chi phí vốn hóa, giao dịch bên liên quan, ý kiến kiểm toán và đối chiếu CFO/LNST/FCF. Với doanh nghiệp tài chính/ngân hàng, các mô hình công nghiệp như Beneish/Jones/REM chỉ dùng tham khảo nếu dữ liệu đủ.",
            ],
            [
                ("Thao túng tài chính - Tổng hợp 4 lớp cảnh báo", financial_manipulation_summary_df),
                ("Thao túng tài chính - Lớp 1 Beneish M-Score", beneish_report_df),
                ("Thao túng tài chính - Lớp 2 Accrual Quality/Sloan", accrual_quality_report_df),
                ("Thao túng tài chính - Lớp 3 Modified Jones/Kothari", modified_jones_report_df),
                ("Thao túng tài chính - Lớp 4 REM - hoạt động thật", rem_report_df),
            ],
            [],
        ),
        # MODULE 3
        ReportSection(
            "3.1 So sánh doanh nghiệp / Tab So sánh doanh nghiệp cùng ngành",
            ["Phần này lấy kết quả từ lần chạy 'So sánh doanh nghiệp' gần nhất trong session. Nếu chưa chạy, bảng sẽ trống."],
            [("Kết quả peer comparison", peer)],
            [(t, f) for t, f in figures if t.startswith("So sánh doanh nghiệp")],
        ),
    ]
    return ReportPackage(title=title, ticker=ticker, sections=sections, warnings=warnings)


def _excel_text_fill(value: object, header: str = "") -> str | None:
    text = str(value or "").lower()
    h = str(header or "").lower()
    if any(k in text for k in ["rủi ro", "cảnh báo", "không đạt", "overvalued", "bán", "bất lợi", "âm"]):
        return BRAND_RED_SOFT
    if any(k in text for k in ["đạt", "tốt", "mạnh", "undervalued", "mua", "thuận lợi", "xanh"]):
        return BRAND_GREEN_SOFT
    if any(k in text for k in ["theo dõi", "trung bình", "cần kiểm tra", "chưa đủ", "mos", "khuyến nghị"]):
        return BRAND_YELLOW_SOFT
    if any(k in h for k in ["tín hiệu", "khuyến nghị", "kết luận", "mức độ", "moat"]):
        return BRAND_YELLOW_SOFT
    return None


def _excel_number_format(header: str, value: object) -> str:
    h = str(header or "").lower()
    if not isinstance(value, (int, float)) or pd.isna(value):
        return "General"
    if any(k in h for k in ["%", "pct", "mos", "roe", "roa", "roic", "wacc", "margin", "tỷ lệ", "điểm"]):
        return '0.0'
    if any(k in h for k in ["giá", "vnd", "đồng", "eps", "oeps", "cp"]):
        return '#,##0'
    if abs(float(value)) >= 100:
        return '#,##0'
    return '0.0'


def _write_dataframe_to_sheet(ws, df: pd.DataFrame, start_row: int, title: str = "") -> int:
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter

    row = start_row
    thin = Side(style="thin", color=BRAND_BORDER)
    box_border = Border(left=thin, right=thin, top=thin, bottom=thin)
    force_numeric_heat = "bảng phân tích sử dụng dòng tiền" in str(title or "").lower()
    if title:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
        cell = ws.cell(row=row, column=1, value=title)
        cell.font = Font(bold=True, size=13, color=BRAND_TEAL_DARK)
        cell.fill = PatternFill("solid", fgColor=BRAND_YELLOW_SOFT)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
        cell.border = box_border
        ws.row_dimensions[row].height = 24
        row += 1
    df = _sanitize_report_df_for_ui(df)
    if df.empty:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
        c = ws.cell(row=row, column=1, value="Chưa có dữ liệu")
        c.fill = PatternFill("solid", fgColor="F8FAFC")
        c.font = Font(italic=True, color="64748B")
        c.alignment = Alignment(wrap_text=True)
        return row + 2
    df = df.replace({pd.NA: None})
    df = df.copy()
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            df[col] = df[col].astype(str)
    headers = [str(c) for c in df.columns]
    header_fill = PatternFill("solid", fgColor=BRAND_TEAL_SOFT)
    for col_idx, header in enumerate(headers, start=1):
        cell = ws.cell(row=row, column=col_idx, value=header)
        cell.font = Font(bold=True, color=BRAND_TEAL_DARK)
        cell.fill = header_fill
        cell.border = box_border
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")
    ws.row_dimensions[row].height = 28
    row += 1
    for _, rec in df.iterrows():
        row_values = rec.tolist()
        for col_idx, value in enumerate(row_values, start=1):
            header = headers[col_idx - 1]
            if isinstance(value, (list, dict, tuple)):
                value = str(value)
            if isinstance(value, float) and pd.isna(value):
                value = None
            cell = ws.cell(row=row, column=col_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = box_border
            cell.number_format = _excel_number_format(header, value)
            fill_color = None
            numeric_value = _numeric_float(value)
            if numeric_value is not None:
                # Heatmap giống app: dương xanh, âm đỏ. Riêng bảng FCF transposed có
                # tiêu đề cột là kỳ/năm nên cần ép heatmap cho mọi ô số ở các cột kỳ.
                if numeric_value < 0:
                    fill_color = BRAND_RED_SOFT
                elif numeric_value > 0 and (force_numeric_heat or any(k in header.lower() for k in ["tăng", "growth", "mos", "điểm", "%", "roe", "roa", "roic", "fcf", "cfo"])):
                    fill_color = BRAND_GREEN_SOFT
            if not fill_color:
                fill_color = _excel_text_fill(value, header)
            if fill_color:
                cell.fill = PatternFill("solid", fgColor=fill_color)
        row += 1
    for col_idx, header in enumerate(headers, start=1):
        series = df.iloc[:, col_idx - 1].astype(str).replace("nan", "") if not df.empty else pd.Series([], dtype=str)
        max_len = max([len(str(header))] + [len(v) for v in series.head(100).tolist()])
        if any(k in str(header).lower() for k in ["nhận xét", "chứng cứ", "nguyên tắc", "mô tả", "ghi chú", "file", "nguồn", "summary"]):
            width = min(max(max_len + 2, 28), 58)
        else:
            width = min(max(max_len + 2, 11), 28)
        ws.column_dimensions[get_column_letter(col_idx)].width = width
    return row + 2


def _add_image_to_sheet(ws, img_bytes: bytes, cell_ref: str, width: int = 800, height: int = 470) -> None:
    from openpyxl.drawing.image import Image as XLImage
    image = XLImage(BytesIO(img_bytes))
    image.width = width
    image.height = height
    ws.add_image(image, cell_ref)


def export_report_xlsx(package: ReportPackage, output_path: str | Path | None = None) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    used: set[str] = set()
    cover = wb.active
    cover.title = _safe_sheet_name("Cover", used)
    cover.sheet_view.showGridLines = False
    cover.column_dimensions["A"].width = 26
    cover.column_dimensions["B"].width = 86
    cover.column_dimensions["C"].width = 22
    cover["A1"] = package.title
    cover.merge_cells("A1:C2")
    cover["A1"].font = Font(bold=True, size=20, color="FFFFFF")
    cover["A1"].fill = PatternFill("solid", fgColor=BRAND_TEAL)
    cover["A1"].alignment = Alignment(wrap_text=True, vertical="center")
    cover["A4"] = "Mục"
    cover["B4"] = "Nội dung"
    for c in ["A4", "B4"]:
        cover[c].font = Font(bold=True, color=BRAND_TEAL_DARK)
        cover[c].fill = PatternFill("solid", fgColor=BRAND_TEAL_SOFT)
    cover["A5"] = "Phạm vi"
    cover["B5"] = "Báo cáo gom toàn bộ nội dung phân tích, bảng dữ liệu, đánh giá, cảnh báo và đồ thị của Tổng quan doanh nghiệp, Định giá chuyên sâu và So sánh doanh nghiệp nếu đã chạy."
    cover["A6"] = "Mã"
    cover["B6"] = package.ticker
    cover["A7"] = "Lưu ý"
    cover["B7"] = "Định dạng màu, card vàng/đỏ/xanh và bảng được mô phỏng theo dashboard app. Biểu đồ được nhúng vào sheet 'Biểu đồ'."
    for row in range(5, 8):
        cover[f"A{row}"].font = Font(bold=True, color=BRAND_TEAL_DARK)
        cover[f"B{row}"].alignment = Alignment(wrap_text=True, vertical="top")
    cover["A9"] = "Các sheet chính"
    cover["A9"].font = Font(bold=True, color=BRAND_TEAL_DARK)
    for i, section in enumerate(package.sections, start=10):
        cover.cell(i, 1, section.title)
        cover.cell(i, 1).fill = PatternFill("solid", fgColor=BRAND_YELLOW_SOFT if i % 2 == 0 else "FFFFFF")
        cover.cell(i, 1).alignment = Alignment(wrap_text=True)
    cover.freeze_panes = "A4"

    # Dedicated chart sheet so charts are immediately visible even if data tables are long.
    all_figures: list[tuple[str, go.Figure]] = []
    for section in package.sections:
        all_figures.extend(section.figures)
    if all_figures:
        ws_chart = wb.create_sheet(_safe_sheet_name("Biểu đồ", used))
        ws_chart.sheet_view.showGridLines = False
        ws_chart["A1"] = "Biểu đồ báo cáo - giống dashboard app"
        ws_chart["A1"].font = Font(bold=True, size=16, color=BRAND_TEAL_DARK)
        ws_chart["A1"].fill = PatternFill("solid", fgColor=BRAND_YELLOW_SOFT)
        ws_chart.column_dimensions["A"].width = 120
        r = 3
        for fig_title, fig in all_figures:
            ws_chart.cell(row=r, column=1, value=fig_title)
            ws_chart.cell(row=r, column=1).font = Font(bold=True, size=12, color=BRAND_TEAL_DARK)
            r += 1
            img_bytes = _fig_to_png_bytes(fig)
            if img_bytes:
                _add_image_to_sheet(ws_chart, img_bytes, f"A{r}", width=850, height=500)
                r += 28
            else:
                ws_chart.cell(row=r, column=1, value="Không xuất được ảnh biểu đồ trên máy này. Hãy chạy install_and_run_app.bat để cài kaleido hoặc dùng fallback matplotlib.")
                r += 2

    for section in package.sections:
        ws = wb.create_sheet(_safe_sheet_name(section.title, used))
        ws.sheet_view.showGridLines = False
        ws.freeze_panes = "A4"
        for col in range(1, 13):
            ws.column_dimensions[chr(64 + col) if col <= 26 else "A"].width = 16
        ws.merge_cells(start_row=1, start_column=1, end_row=2, end_column=8)
        ws["A1"] = section.title
        ws["A1"].font = Font(bold=True, size=16, color="FFFFFF")
        ws["A1"].fill = PatternFill("solid", fgColor=BRAND_TEAL)
        ws["A1"].alignment = Alignment(wrap_text=True, vertical="center")
        row = 4
        # Paragraphs as soft cards.
        for para in section.paragraphs:
            text = _clean_text(para)
            if not text:
                continue
            ws.merge_cells(start_row=row, start_column=1, end_row=row + 1, end_column=8)
            c = ws.cell(row=row, column=1, value=text)
            c.alignment = Alignment(wrap_text=True, vertical="top")
            c.fill = PatternFill("solid", fgColor=BRAND_YELLOW_SOFT if any(k in text.lower() for k in ["khuyến nghị", "mos", "phân loại", "cảnh báo"]) else "F8FFFB")
            c.font = Font(color=BRAND_TEAL_DARK if "khuyến nghị" not in text.lower() else "7A4A00", bold="khuyến nghị" in text.lower())
            ws.row_dimensions[row].height = 34
            ws.row_dimensions[row + 1].height = 34
            row += 3
        # Put section charts before long data tables.
        for fig_title, fig in section.figures:
            img_bytes = _fig_to_png_bytes(fig)
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=8)
            ws.cell(row=row, column=1, value=fig_title)
            ws.cell(row=row, column=1).font = Font(bold=True, size=13, color=BRAND_TEAL_DARK)
            ws.cell(row=row, column=1).fill = PatternFill("solid", fgColor=BRAND_YELLOW_SOFT)
            row += 1
            if img_bytes:
                _add_image_to_sheet(ws, img_bytes, f"A{row}", width=800, height=470)
                row += 27
            else:
                ws.cell(row=row, column=1, value=f"Không xuất được ảnh biểu đồ: {fig_title}. Cài/kiểm tra kaleido hoặc matplotlib nếu cần xuất ảnh chart.")
                row += 2
        for table_title, table_df in section.tables:
            row = _write_dataframe_to_sheet(ws, table_df, row, table_title)
    buf = BytesIO()
    wb.save(buf)
    data = buf.getvalue()
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_bytes(data)
    return data


def _docx_set_cell_shading(cell, fill: str) -> None:
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)
    except Exception:
        pass


def _docx_set_cell_text(cell, text: object, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = 0
    run = p.add_run("" if text is None or (isinstance(text, float) and pd.isna(text)) else str(text))
    run.bold = bold
    if color:
        try:
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor.from_string(color)
        except Exception:
            pass


def _docx_add_card(doc, text: str, fill: str = "F8FFFB", border_color: str = BRAND_TEAL_SOFT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _docx_set_cell_shading(cell, fill)
    _docx_set_cell_text(cell, text, bold=False, color=BRAND_TEAL_DARK)
    doc.add_paragraph("")


def _docx_add_table(doc, df: pd.DataFrame, title: str) -> None:
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    df = _sanitize_report_df_for_ui(df)
    heading = doc.add_heading(title, level=3)
    for run in heading.runs:
        run.font.color.rgb = None
    if df.empty:
        doc.add_paragraph("Chưa có dữ liệu.")
        return
    # Giữ rộng vừa phải để Word không vỡ layout; dữ liệu đầy đủ vẫn ở Excel.
    cols = list(df.columns)[:10]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, col in enumerate(cols):
        cell = table.rows[0].cells[i]
        _docx_set_cell_shading(cell, BRAND_TEAL_SOFT)
        _docx_set_cell_text(cell, str(col), bold=True, color=BRAND_TEAL_DARK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for _, rec in df[cols].iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(rec.tolist()):
            header = str(cols[i])
            cell = row_cells[i]
            text = "" if pd.isna(value) else str(value)
            _docx_set_cell_text(cell, text, bold=False, color="123D3A")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            fill = None
            if isinstance(value, (int, float)) and not pd.isna(value):
                if float(value) < 0:
                    fill = BRAND_RED_SOFT
                elif float(value) > 0 and any(k in header.lower() for k in ["tăng", "growth", "mos", "điểm", "%", "roe", "roa", "roic", "fcf", "cfo"]):
                    fill = BRAND_GREEN_SOFT
            if not fill:
                fill = _excel_text_fill(value, header)
            if fill:
                _docx_set_cell_shading(cell, fill)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
    doc.add_paragraph("")


def export_report_docx(package: ReportPackage, output_path: str | Path | None = None) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.section import WD_ORIENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(package.title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(BRAND_TEAL_DARK)
    _docx_add_card(doc, "Báo cáo tự động gồm toàn bộ nội dung phân tích, bảng dữ liệu, đánh giá, cảnh báo và đồ thị của các phần trong app. Định dạng màu mô phỏng dashboard Trecapital.", fill=BRAND_YELLOW_SOFT)

    for sec in package.sections:
        h = doc.add_heading(sec.title, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(BRAND_TEAL_DARK)
        for para in sec.paragraphs:
            text = _clean_text(para)
            if text:
                fill = BRAND_YELLOW_SOFT if any(k in text.lower() for k in ["khuyến nghị", "mos", "phân loại", "cảnh báo"]) else "F8FFFB"
                _docx_add_card(doc, text, fill=fill)
        # Chart trước bảng để người đọc thấy ngay như app.
        for fig_title, fig in sec.figures:
            h3 = doc.add_heading(fig_title, level=3)
            for run in h3.runs:
                run.font.color.rgb = RGBColor.from_string(BRAND_TEAL_DARK)
            img_bytes = _fig_to_png_bytes(fig)
            if img_bytes:
                doc.add_picture(BytesIO(img_bytes), width=Inches(9.6))
            else:
                _docx_add_card(doc, "Không xuất được ảnh biểu đồ. Hãy chạy install_and_run_app.bat để cài kaleido hoặc dùng fallback matplotlib.", fill=BRAND_RED_SOFT)
        for table_title, table_df in sec.tables:
            _docx_add_table(doc, table_df, table_title)
    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_bytes(data)
    return data


def _pdf_table_data(df: pd.DataFrame, max_cols: int = 8, max_rows: int = 28) -> list[list[str]]:
    df = _sanitize_report_df_for_ui(df)
    if df.empty:
        return [["Chưa có dữ liệu"]]
    cols = list(df.columns)[:max_cols]
    out = [[str(c) for c in cols]]
    for _, rec in df[cols].head(max_rows).iterrows():
        out.append(["" if pd.isna(v) else str(v)[:180] for v in rec.tolist()])
    if len(df) > max_rows:
        out.append([f"... còn {len(df) - max_rows} dòng; xem đầy đủ trong file Excel/Word"] + [""] * (len(cols) - 1))
    return out


def _pdf_escape_para(text: str) -> str:
    return _clean_text(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>")


def _export_report_pdf_reportlab(package: ReportPackage, output_path: str | Path | None = None) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether

    font_regular = "Helvetica"
    font_bold = "Helvetica-Bold"
    for font_path, font_name in [
        ("/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf", "NotoSans"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "DejaVuSans"),
        ("C:/Windows/Fonts/arial.ttf", "Arial"),
    ]:
        try:
            if Path(font_path).exists():
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                font_regular = font_name
                break
        except Exception:
            pass
    for font_path, font_name in [
        ("/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf", "NotoSans-Bold"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", "DejaVuSans-Bold"),
        ("C:/Windows/Fonts/arialbd.ttf", "Arial-Bold"),
    ]:
        try:
            if Path(font_path).exists():
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                font_bold = font_name
                break
        except Exception:
            pass

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=0.38*inch, rightMargin=0.38*inch, topMargin=0.38*inch, bottomMargin=0.38*inch)
    styles = getSampleStyleSheet()
    styles["Title"].fontName = font_bold
    styles["Title"].textColor = colors.HexColor("#" + BRAND_TEAL_DARK)
    styles["Heading3"].fontName = font_bold
    styles["Heading3"].textColor = colors.HexColor("#" + BRAND_TEAL_DARK)
    styles.add(ParagraphStyle(name="SmallVN", parent=styles["BodyText"], fontName=font_regular, fontSize=8.2, leading=10.2, alignment=TA_LEFT))
    styles.add(ParagraphStyle(name="CardVN", parent=styles["SmallVN"], backColor=colors.HexColor("#F8FFFB"), borderColor=colors.HexColor("#" + BRAND_TEAL_SOFT), borderWidth=0.5, borderPadding=5))
    styles.add(ParagraphStyle(name="YellowCardVN", parent=styles["SmallVN"], textColor=colors.HexColor("#5F3B00"), backColor=colors.HexColor("#" + BRAND_YELLOW_SOFT), borderColor=colors.HexColor("#" + BRAND_YELLOW), borderWidth=0.8, borderPadding=5))
    styles.add(ParagraphStyle(name="HeadingTeal", parent=styles["Heading1"], fontName=font_bold, textColor=colors.HexColor("#" + BRAND_TEAL_DARK), fontSize=16, leading=20))
    story = [Paragraph(_pdf_escape_para(package.title), styles["Title"]), Spacer(1, 8)]
    story.append(Paragraph("Báo cáo tự động gồm nội dung phân tích, bảng đánh giá và đồ thị của các phần trong app.", styles["YellowCardVN"]))
    story.append(Spacer(1, 12))
    for sec_idx, sec in enumerate(package.sections):
        if sec_idx > 0:
            story.append(PageBreak())
        story.append(Paragraph(_pdf_escape_para(sec.title), styles["HeadingTeal"]))
        for para in sec.paragraphs:
            text = _clean_text(para)
            if text:
                style = styles["YellowCardVN"] if any(k in text.lower() for k in ["khuyến nghị", "mos", "phân loại", "cảnh báo"]) else styles["CardVN"]
                story.append(Paragraph(_pdf_escape_para(text), style))
                story.append(Spacer(1, 5))
        for fig_title, fig in sec.figures:
            img_bytes = _fig_to_png_bytes(fig, width=1100, height=620, scale=2)
            blocks = [Paragraph(_pdf_escape_para(fig_title), styles["Heading3"])]
            if img_bytes:
                blocks.append(Image(BytesIO(img_bytes), width=8.8*inch, height=4.95*inch))
            else:
                blocks.append(Paragraph("Không xuất được ảnh biểu đồ. Hãy chạy install_and_run_app.bat để cài kaleido hoặc dùng fallback matplotlib.", styles["SmallVN"]))
            story.append(KeepTogether(blocks))
            story.append(Spacer(1, 8))
        for table_title, table_df in sec.tables:
            story.append(Paragraph(_pdf_escape_para(table_title), styles["Heading3"]))
            data = _pdf_table_data(table_df)
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#" + BRAND_TEAL_SOFT)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#" + BRAND_TEAL_DARK)),
                ("FONTNAME", (0, 0), (-1, 0), font_bold),
                ("FONTNAME", (0, 1), (-1, -1), font_regular),
                ("FONTSIZE", (0, 0), (-1, -1), 6.3),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#" + BRAND_BORDER)),
            ]))
            story.append(table)
            story.append(Spacer(1, 8))
    doc.build(story)
    data = buf.getvalue()
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_bytes(data)
    return data


def _pil_font(size: int = 18, bold: bool = False):
    from PIL import ImageFont
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for fp in candidates:
        try:
            if Path(fp).exists():
                return ImageFont.truetype(fp, size=size)
        except Exception:
            pass
    return ImageFont.load_default()


def _draw_wrapped(draw, xy: tuple[int, int], text: str, font, fill: tuple[int, int, int], max_width: int, line_spacing: int = 6, max_lines: int | None = None) -> int:
    x, y = xy
    words = str(text or "").replace("\n", " \n ").split()
    lines: list[str] = []
    cur = ""
    for word in words:
        if word == "\n":
            lines.append(cur)
            cur = ""
            continue
        test = word if not cur else f"{cur} {word}"
        try:
            w = draw.textbbox((0, 0), test, font=font)[2]
        except Exception:
            w = len(test) * 8
        if w <= max_width or not cur:
            cur = test
        else:
            lines.append(cur)
            cur = word
    if cur:
        lines.append(cur)
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = lines[-1][:max(0, len(lines[-1])-1)] + "…"
    line_h = int(getattr(font, "size", 16) * 1.28) + line_spacing
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h
    return y


def _export_report_pdf_pil(package: ReportPackage, output_path: str | Path | None = None) -> bytes:
    # Dependency-light PDF fallback: render report pages as images and save as PDF.
    from PIL import Image, ImageDraw
    W, H = 1654, 1169  # A4 landscape at ~140 dpi.
    margin = 58
    teal = tuple(int(BRAND_TEAL[i:i+2], 16) for i in (0, 2, 4))
    teal_dark = tuple(int(BRAND_TEAL_DARK[i:i+2], 16) for i in (0, 2, 4))
    yellow_soft = tuple(int(BRAND_YELLOW_SOFT[i:i+2], 16) for i in (0, 2, 4))
    border = tuple(int(BRAND_BORDER[i:i+2], 16) for i in (0, 2, 4))
    text_col = (18, 52, 59)
    pages = []

    def new_page(title: str = ""):
        img = Image.new("RGB", (W, H), "white")
        d = ImageDraw.Draw(img)
        d.rectangle([0, 0, W, 24], fill=teal)
        if title:
            d.text((margin, 42), title, font=_pil_font(26, bold=True), fill=teal_dark)
        return img, d, 86 if title else margin

    img, d, y = new_page(package.title)
    d.rounded_rectangle([margin, y, W - margin, y + 92], radius=22, fill=yellow_soft, outline=border, width=2)
    y = _draw_wrapped(d, (margin + 22, y + 16), "Báo cáo tự động gồm nội dung phân tích, bảng đánh giá, cảnh báo và đồ thị. PDF fallback không cần reportlab nên sẽ không lỗi thiếu thư viện.", _pil_font(20), text_col, W - 2 * margin - 44)
    pages.append(img)

    for sec in package.sections:
        img, d, y = new_page(sec.title)
        for para in sec.paragraphs:
            text = _clean_text(para)
            if not text:
                continue
            box_h = 116
            if y + box_h > H - margin:
                pages.append(img)
                img, d, y = new_page(sec.title)
            fill = yellow_soft if any(k in text.lower() for k in ["khuyến nghị", "mos", "phân loại", "cảnh báo"]) else (248, 255, 251)
            d.rounded_rectangle([margin, y, W - margin, y + box_h], radius=18, fill=fill, outline=border, width=1)
            _draw_wrapped(d, (margin + 18, y + 12), text, _pil_font(16), text_col, W - 2 * margin - 36, max_lines=5)
            y += box_h + 18
        for fig_title, fig in sec.figures:
            if y + 610 > H - margin:
                pages.append(img)
                img, d, y = new_page(sec.title)
            d.text((margin, y), fig_title, font=_pil_font(20, bold=True), fill=teal_dark)
            y += 34
            img_bytes = _fig_to_png_bytes(fig, width=1200, height=650)
            if img_bytes:
                try:
                    chart = Image.open(BytesIO(img_bytes)).convert("RGB")
                    chart.thumbnail((W - 2 * margin, 520))
                    d.rectangle([margin, y, margin + chart.width + 8, y + chart.height + 8], outline=border, width=1)
                    img.paste(chart, (margin + 4, y + 4))
                    y += chart.height + 26
                except Exception:
                    d.text((margin, y), "Không xuất được ảnh biểu đồ.", font=_pil_font(16), fill=(185, 28, 28))
                    y += 36
            else:
                d.text((margin, y), "Không xuất được ảnh biểu đồ.", font=_pil_font(16), fill=(185, 28, 28))
                y += 36
        for table_title, table_df in sec.tables:
            df = _sanitize_report_df_for_ui(table_df)
            if y + 180 > H - margin:
                pages.append(img)
                img, d, y = new_page(sec.title)
            d.text((margin, y), table_title, font=_pil_font(20, bold=True), fill=teal_dark)
            y += 34
            if df.empty:
                d.text((margin, y), "Chưa có dữ liệu", font=_pil_font(15), fill=(100, 116, 139))
                y += 28
                continue
            cols = list(df.columns)[:5]
            col_w = (W - 2 * margin) // max(len(cols), 1)
            row_h = 42
            d.rectangle([margin, y, W - margin, y + row_h], fill=(234, 247, 241), outline=border)
            for i, c in enumerate(cols):
                d.text((margin + i * col_w + 6, y + 10), textwrap.shorten(str(c), width=24, placeholder="…"), font=_pil_font(13, bold=True), fill=teal_dark)
            y += row_h
            for _, rec in df[cols].head(10).iterrows():
                if y + row_h > H - margin:
                    pages.append(img)
                    img, d, y = new_page(sec.title)
                d.rectangle([margin, y, W - margin, y + row_h], fill="white", outline=border)
                for i, v in enumerate(rec.tolist()):
                    d.text((margin + i * col_w + 6, y + 9), textwrap.shorten("" if pd.isna(v) else str(v), width=26, placeholder="…"), font=_pil_font(12), fill=text_col)
                y += row_h
            y += 18
        pages.append(img)
    buf = BytesIO()
    if not pages:
        img, _, _ = new_page(package.title)
        pages = [img]
    pages[0].save(buf, format="PDF", save_all=True, append_images=pages[1:])
    data = buf.getvalue()
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        Path(output_path).write_bytes(data)
    return data


def export_report_pdf(package: ReportPackage, output_path: str | Path | None = None) -> bytes:
    """Export PDF without crashing when reportlab is absent.

    Primary path: reportlab (vector text/tables). Fallback: Pillow image-PDF, which is enough
    for review/printing and avoids ModuleNotFoundError on machines that haven't reinstalled deps.
    """
    try:
        return _export_report_pdf_reportlab(package, output_path)
    except ModuleNotFoundError:
        return _export_report_pdf_pil(package, output_path)
    except ImportError:
        return _export_report_pdf_pil(package, output_path)





def _escape_html_text(value: object) -> str:
    return html.escape(str(value or ""))



def _num_or_none(value: object) -> float | None:
    """Parse a value that may already be numeric or may be a formatted string."""
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    if isinstance(value, (int, float)):
        num = float(value)
        return num if math.isfinite(num) else None
    raw = str(value).strip()
    if not raw or raw.lower() in {"nan", "none", "n/a", "na", "null"}:
        return None
    raw = raw.replace("%", "").replace(" ", "").replace(",", "")
    try:
        num = float(raw)
    except Exception:
        return None
    return num if math.isfinite(num) else None


def _fmt_bil0(value: object) -> str:
    num = _num_or_none(value)
    return "N/A" if num is None else f"{num:,.0f} tỷ đồng"


def _fmt_vnd0(value: object) -> str:
    num = _num_or_none(value)
    return "N/A" if num is None else f"{num:,.0f} đ/cp"


def _fmt_pct1(value: object) -> str:
    num = _num_or_none(value)
    return "N/A" if num is None else f"{num:,.1f}%"


def _fmt_ratio1(value: object) -> str:
    num = _num_or_none(value)
    return "N/A" if num is None else f"{num:,.1f}x"


def _fmt_score1(value: object, suffix: str = "/100") -> str:
    num = _num_or_none(value)
    return "N/A" if num is None else f"{num:,.1f}{suffix}"


def _fmt_text_na(value: object) -> str:
    if value is None:
        return "N/A"
    try:
        if pd.isna(value):
            return "N/A"
    except Exception:
        pass
    text = str(value).strip()
    return "N/A" if not text or text.lower() in {"nan", "none", "n/a", "na", "null"} else text


def _latest_any(row: dict[str, object], keys: list[str]) -> object:
    for key in keys:
        if key in row:
            value = row.get(key)
            if _num_or_none(value) is not None:
                return value
    return None


def _ratio_from_latest(row: dict[str, object], numerator_keys: list[str], denominator_keys: list[str]) -> float | None:
    num = _num_or_none(_latest_any(row, numerator_keys))
    den = _num_or_none(_latest_any(row, denominator_keys))
    if num is None or den is None or abs(den) < 1e-12:
        return None
    return num / den


def _format_number_for_source_rule(value: object, header: str = "") -> str:
    """Format numbers for the printable report following the project source rule.

    - Tỷ đồng/VND/số lượng: 0 decimals.
    - Phần trăm: 1 decimal.
    - Hệ số/ratio/score: 1 decimal.
    - Text/date columns: keep text.
    """
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    h = str(header or "").strip().lower()
    raw = str(value).strip()
    if raw == "":
        return ""
    # Keep label/text columns untouched.
    label_keys = [
        "period", "kỳ", "năm", "quý", "mã", "ticker", "sàn", "ngành", "nguồn", "đường dẫn",
        "file", "url", "công thức", "diễn giải", "nhận xét", "nội dung", "câu hỏi", "lý do",
        "cơ sở", "tín hiệu", "mức độ", "khuyến nghị", "kết luận", "loại", "vai trò",
        "beta_source", "wacc_formula_detail", "tiêu chí", "chỉ tiêu", "phương pháp", "kịch bản",
    ]
    if any(k in h for k in label_keys):
        return raw
    cleaned = raw.replace(" ", "").replace(",", "")
    suffix_percent = False
    if cleaned.endswith("%"):
        suffix_percent = True
        cleaned = cleaned[:-1]
    try:
        num = float(cleaned)
    except Exception:
        return raw
    if not math.isfinite(num):
        return ""

    percent_keys = ["%", "pct", "tỷ lệ", "ty le", "biên", "bien", "margin", "growth", "tăng trưởng", "tang truong", "mos", "roe", "roa", "roic", "wacc", "yield"]
    ratio_keys = ["p/e", "p/b", "p/s", "ev/", "hệ số", "he so", "vòng quay", "vong quay", "turnover", "multiplier", "ccc", "d/e", "nợ vay/", "no vay/", "debt/", "cfo/", "fcf/", "beta"]
    score_keys = ["điểm", "diem", "score", "trọng số", "trong so"]
    price_keys = ["giá", "gia", "vnd", "đồng", "dong", "eps", "oeps", "bvps", "market cap", "vốn hóa", "von hoa"]
    bil_keys = ["_bil", "tỷ đồng", "ty dong", "doanh thu", "lợi nhuận", "loi nhuan", "lntt", "lnst", "ebit", "ebitda", "cfo", "fcf", "capex", "cash", "tiền", "tien", "nợ", "no", "vốn", "von", "tài sản", "tai san", "equity", "debt", "inventory", "working capital", "capital", "investment"]

    # V23.68: Columns such as "Giá mua MOS 30%" contain both "giá" and "MOS/%".
    # They are price-per-share thresholds (đ/cp), not percentage metrics. Give price
    # semantics precedence over percentage semantics for all MOS buy-price headers.
    mos_price_header = ("giá" in h or "gia" in h) and "mos" in h

    if mos_price_header:
        return f"{num:,.0f}"
    if suffix_percent or any(k in h for k in percent_keys):
        return f"{num:,.1f}%"
    if any(k in h for k in ratio_keys):
        return f"{num:,.1f}"
    if any(k in h for k in score_keys):
        return f"{num:,.1f}"
    if any(k in h for k in price_keys):
        return f"{num:,.0f}"
    if any(k in h for k in bil_keys):
        return f"{num:,.0f}"
    if abs(num - round(num)) < 1e-9:
        return f"{num:,.0f}"
    return f"{num:,.1f}"


def _format_print_df(df: pd.DataFrame, max_rows: int | None = None) -> pd.DataFrame:
    out = _df(df).copy()
    if max_rows is not None and len(out) > max_rows:
        out = out.head(max_rows)
    formatted = pd.DataFrame(index=out.index)
    for col in out.columns:
        formatted[col] = out[col].map(lambda v, c=str(col): _format_number_for_source_rule(v, c))
    return formatted


def _render_report_card_html(title: str, body: str, *, tone: str = "normal") -> str:
    if tone == "yellow":
        bg = "linear-gradient(135deg,#FFF3C4 0%,#FFFFFF 100%)"
        border = "rgba(245,178,27,.65)"
    elif tone == "green":
        bg = "linear-gradient(135deg,#EAF7F1 0%,#FFFFFF 100%)"
        border = "rgba(11,127,117,.35)"
    else:
        bg = "#FFFFFF"
        border = "rgba(11,127,117,.22)"
    return f"""
    <div class="tre-report-card" style="border:1.4px solid {border}; border-radius:18px; padding:14px 16px; margin:10px 0 14px 0; background:{bg}; box-shadow:0 8px 24px rgba(15,23,42,.06); break-inside:avoid; page-break-inside:avoid;">
      <div style="font-size:15px; font-weight:900; color:#064E47; margin-bottom:6px;">{_escape_html_text(title)}</div>
      <div style="font-size:13px; line-height:1.55; color:#0F172A; white-space:pre-wrap;">{_escape_html_text(body)}</div>
    </div>
    """




def _is_numeric_like(value: object) -> bool:
    try:
        if value is None or (isinstance(value, float) and math.isnan(value)):
            return False
        text = str(value).strip().replace(",", "").replace("%", "")
        if text in {"", "-", "N/A", "nan", "None"}:
            return False
        float(text)
        return True
    except Exception:
        return False


def _numeric_float(value: object) -> float | None:
    try:
        if value is None or (isinstance(value, float) and math.isnan(value)):
            return None
        text = str(value).strip().replace(",", "").replace("%", "")
        if text in {"", "-", "N/A", "nan", "None"}:
            return None
        return float(text)
    except Exception:
        return None


def _cell_heat_class(value: object, header: str = "") -> str:
    """Return app-like heat class for a print-table cell."""
    h = str(header or "").lower()
    text = str(value or "").lower()
    num = _numeric_float(value)
    if any(k in text for k in ["rủi ro", "cảnh báo", "không đạt", "yếu", "xấu", "âm", "overvalued", "bán"]):
        return "heat-red"
    if any(k in text for k in ["tốt", "mạnh", "đạt", "undervalued", "mua", "thuận lợi", "xanh"]):
        return "heat-green"
    if any(k in text for k in ["theo dõi", "trung bình", "cần kiểm tra", "chưa đủ", "mos"]):
        return "heat-yellow"
    if num is None:
        return ""
    if any(k in h for k in ["điểm", "score", "tỷ lệ đạt", "điểm nhiệt"]):
        if num >= 80:
            return "heat-green-strong"
        if num >= 65:
            return "heat-green"
        if num >= 45:
            return "heat-yellow"
        if num >= 25:
            return "heat-orange"
        return "heat-red"
    if any(k in h for k in ["mos", "tăng trưởng", "%", "cfo/l", "fcf/l", "roic", "roe", "roa", "wacc", "biên"]):
        if num > 0:
            return "num-pos"
        if num < 0:
            return "num-neg"
    return ""



def _sanitize_report_df_for_ui(df: pd.DataFrame) -> pd.DataFrame:
    src = _df(df)
    if src.empty:
        return src
    hidden = {"source", "Source", "Nguồn", "Nguồn/URL", "Nguồn dữ liệu", "Nguồn/logic", "Nguồn / logic", "URL", "Truy vấn", "updated_at", "note", "raw_path", "File nguồn/cache", "beta_source"}
    out = src.drop(columns=[c for c in src.columns if str(c) in hidden or "url" in str(c).lower() or str(c).lower() == "source"], errors="ignore").copy()
    replacements = {
        "FireAnt": "Dữ liệu ưu tiên",
        "Vietstock": "Dữ liệu ưu tiên",
        "Simplize": "Danh sách cùng ngành",
        "KBS": "nhóm trực tuyến",
        "VCI": "nhóm trực tuyến",
        "CafeF": "Tham khảo",
        "raw_data": "nhật ký nội bộ",
        "data_cache": "bộ nhớ dữ liệu",
    }
    for c in out.columns:
        if out[c].dtype == object:
            out[c] = out[c].map(lambda v: _replace_terms(str(v), replacements) if v is not None else v)
    return out


def _replace_terms(text: str, replacements: dict[str, str]) -> str:
    for raw, public in replacements.items():
        text = text.replace(raw, public)
    text = re.sub(r"https?://\S+", "liên kết nội bộ", text, flags=re.I)
    text = re.sub(r"/?[^\s<>]*\.(?:csv|json|html|txt|xlsm|xlsx|md)", "file nội bộ", text, flags=re.I)
    text = re.sub(r"/(?:mnt|home|Users|raw_data|data_cache)[^\s<>]+", "đường dẫn nội bộ", text)
    return text


def _render_static_print_table_html(df: pd.DataFrame, title: str = "", max_rows: int | None = None) -> str:
    """Render a non-scrollable HTML table so browser print/Save as PDF prints every row.

    This helper is intentionally used only on the consolidated report page. It does not replace
    interactive st.dataframe tables in Tổng quan doanh nghiệp/2/3, so analytical screens and calculations remain unchanged.
    """
    src = _sanitize_report_df_for_ui(df)
    if src.empty:
        return ""
    truncated_note = ""
    if max_rows is not None and len(src) > max_rows:
        truncated_note = f"<div class='tre-table-footnote'>Đang hiển thị {int(max_rows)}/{len(src)} dòng theo chế độ gọn để in nhanh.</div>"
        src = src.head(max_rows)
    try:
        src = _format_print_df(src, None)
    except Exception:
        pass
    # Keep all columns, but put the table in a screen-only horizontal wrapper; print mode expands it.
    headers = [str(c) for c in src.columns]

    def _col_class(col_name: str) -> str:
        c = str(col_name or "").strip().lower()
        classes = []
        if c in {"mã", "ma", "ticker", "symbol"}:
            classes.append("col-code")
        if c in {"kỳ", "ky", "period", "kỳ mới nhất", "ky moi nhat"} or c.startswith("kỳ ") or c.startswith("ky "):
            classes.append("col-period")
        if any(k in c for k in ["diễn giải", "nhận xét", "nội dung", "câu hỏi", "nguyên tắc", "cơ sở", "công thức", "bằng chứng"]):
            classes.append("col-long-text")
        return " ".join(classes)

    title_l = str(title or "").lower()
    table_classes = ["tre-static-table"]
    block_classes = ["tre-static-table-block"]
    if "roic" in title_l and "đầu tư" in title_l:
        table_classes.append("tre-table-roic")
        block_classes.append("tre-block-roic")
    if "peer" in title_l or "so sánh" in title_l:
        table_classes.append("tre-table-peer")
        block_classes.append("tre-block-peer")

    thead = "".join(f"<th class='{_col_class(h)}'>{html.escape(h)}</th>" for h in headers)
    rows = []
    for _, row in src.iterrows():
        cells = []
        for col in headers:
            value = row.get(col, "")
            text = "" if pd.isna(value) else str(value)
            cls = " ".join([_cell_heat_class(text, col), _col_class(col)]).strip()
            cells.append(f"<td class='{cls}'>{html.escape(text)}</td>")
        rows.append("<tr>" + "".join(cells) + "</tr>")
    return f"""
    <div class="{' '.join(block_classes)}">
      {f'<div class="tre-report-table-title">📋 {html.escape(str(title))}</div>' if title else ''}
      <div class="tre-static-table-wrap">
        <table class="{' '.join(table_classes)}">
          <thead><tr>{thead}</tr></thead>
          <tbody>{''.join(rows)}</tbody>
        </table>
      </div>
      {truncated_note}
    </div>
    """


def _table_is_evaluation_like(title: str, df: pd.DataFrame) -> bool:
    t = str(title or "").lower()
    cols = " ".join([str(c).lower() for c in _df(df).columns])
    keywords = [
        "đánh giá", "chấm điểm", "score", "tiêu chí", "porter", "moat", "chuỗi giá trị",
        "định giá", "dải giá trị", "phương pháp", "cảnh báo", "rủi ro", "mos", "tín hiệu", "thao túng", "beneish", "accrual", "jones", "rem",
    ]
    return any(k in t for k in keywords) or any(k in cols for k in ["nhận xét", "diễn giải", "tín hiệu", "đánh giá", "điểm", "kết luận"])


def _row_note_for_static_report(title: str, rowd: dict[str, Any]) -> str:
    """Create visible row notes for the print report from the same row-level information used in app tables."""
    title_l = str(title or "").lower()
    parts: list[str] = []
    # Strategic assessment table.
    if any(k in rowd for k in ["Nội dung cần đánh giá", "Câu hỏi đánh giá"]):
        q = rowd.get("Nội dung cần đánh giá", rowd.get("Câu hỏi đánh giá", "Đánh giá trọng yếu"))
        parts.append(f"{q}")
        parts.append(f"Kết luận: {rowd.get('Kết luận theo mã', 'N/A')}")
        parts.append(f"Chứng cứ: {rowd.get('Số liệu/chứng cứ chính', 'N/A')}")
        parts.append(f"Nguyên tắc: {rowd.get('Nguyên tắc áp dụng riêng', 'N/A')}")
        return "\n".join([str(p) for p in parts if str(p).strip()])
    # Valuation method table.
    if "Phương pháp" in rowd:
        parts.append(f"Phương pháp: {rowd.get('Phương pháp', 'N/A')}")
        for c in ["Giá trị/cp", "Giá trị hợp lý/cp", "Fair value", "Base value", "Trọng số", "Trọng số %", "Giá trị weighted", "MOS %", "Tín hiệu"]:
            if c in rowd:
                parts.append(f"{c}: {rowd.get(c)}")
        for c in ["Công thức", "Diễn giải", "Ghi chú", "Nguồn", "Lý do dùng", "Nguyên tắc"]:
            if c in rowd and str(rowd.get(c, "")).strip():
                parts.append(f"{c}: {rowd.get(c)}")
        if len(parts) == 1:
            parts.append("Ghi chú: phương pháp này được app đưa vào dải giá trị nội tại nếu có đủ dữ liệu và trọng số phù hợp với loại doanh nghiệp.")
        return "\n".join(parts)
    # Value range.
    if "Dải giá trị" in title or "dải giá trị" in title_l or ("Chỉ tiêu" in rowd and "Giá trị/cp" in rowd):
        metric = rowd.get("Chỉ tiêu", "Chỉ tiêu định giá")
        val = rowd.get("Giá trị/cp", rowd.get("Giá trị", "N/A"))
        return f"{metric}: {val}\nÝ nghĩa: đây là một mốc trong dải giá trị nội tại/MOS. App không dùng một fair value duy nhất mà trình bày low-base-high-weighted và giá mua theo MOS yêu cầu."
    # Porter / moat.
    if "Nhóm Porter/Moat" in rowd:
        parts.append(f"Nhóm Porter/Moat: {rowd.get('Nhóm Porter/Moat')}")
        for c in ["Điểm đạt", "Trọng số %", "Tỷ lệ đạt %", "Tín hiệu", "Đánh giá sơ bộ", "Nhận xét tự động", "Diễn giải", "Bằng chứng/cần kiểm tra"]:
            if c in rowd and str(rowd.get(c, "")).strip():
                parts.append(f"{c}: {rowd.get(c)}")
        return "\n".join(parts)
    # Value chain.
    if "Hoạt động chuỗi giá trị" in rowd:
        parts.append(f"Hoạt động chuỗi giá trị: {rowd.get('Hoạt động chuỗi giá trị')}")
        for c in ["Điểm nhiệt", "Mức độ", "Tín hiệu", "Đánh giá sơ bộ", "Tác động", "Bằng chứng hiện có/cần tìm", "Nhận xét tự động", "Diễn giải"]:
            if c in rowd and str(rowd.get(c, "")).strip():
                parts.append(f"{c}: {rowd.get(c)}")
        return "\n".join(parts)
    # Scenario/risk.
    if "Kịch bản" in rowd or "rủi ro" in title_l:
        parts.append(f"Kịch bản/rủi ro: {rowd.get('Kịch bản', rowd.get('Rủi ro', 'N/A'))}")
        for c in ["Tác động", "Xác suất", "Mức độ", "Cảnh báo", "Biện pháp", "Ghi chú", "Nhận xét"]:
            if c in rowd and str(rowd.get(c, "")).strip():
                parts.append(f"{c}: {rowd.get(c)}")
        return "\n".join(parts)
    # Generic scorecards / alerts.
    important_cols = [c for c in ["Nhóm tiêu chí", "Tiêu chí", "Chỉ tiêu", "Điểm", "Trọng số", "Tỷ lệ đạt", "Tín hiệu", "Nhận xét tự động", "Diễn giải", "Đánh giá", "Mức độ", "Cảnh báo"] if c in rowd]
    if important_cols:
        for c in important_cols:
            v = rowd.get(c)
            if str(v).strip() and str(v).lower() != "nan":
                parts.append(f"{c}: {v}")
        return "\n".join(parts)
    return ""


def _render_static_table_notes_html(title: str, df: pd.DataFrame, max_notes: int = 80) -> str:
    src = _df(df)
    if src.empty or not _table_is_evaluation_like(title, src):
        return ""
    notes = []
    for i, (_, row) in enumerate(src.iterrows(), start=1):
        if i > max_notes:
            notes.append(f"<div class='tre-table-note-card muted'>Còn {len(src)-max_notes} dòng note khác. Chọn chế độ Đầy đủ nếu cần kiểm tra toàn bộ bảng trong app.</div>")
            break
        note = _row_note_for_static_report(title, row.to_dict())
        if not note.strip():
            continue
        notes.append(
            f"<div class='tre-table-note-card'><b>Note dòng {i}:</b><br>{html.escape(note).replace(chr(10), '<br>')}</div>"
        )
    if not notes:
        return ""
    return f"""
    <div class='tre-table-notes-block'>
      <div class='tre-table-notes-title'>📝 Note / nhận xét / diễn giải đi kèm bảng: {html.escape(str(title))}</div>
      {''.join(notes)}
    </div>
    """

def render_report_package_as_app_page(
    package: ReportPackage,
    *,
    show_export_hint: bool = True,
    table_height: int = 380,
    max_rows_per_table: int | None = None,
) -> None:
    """Render the full report in Streamlit using app-like cards/tables/Plotly charts.

    This page is intended for browser printing. It renders all package sections in a single long page
    so Windows/Chrome can Save as PDF while preserving dashboard colors, cards, tables and charts.
    """
    import streamlit as st

    _inject_print_page_css(st)
    st.markdown(
        """
        <style>
        .tre-report-title-card {
            border: 1.6px solid rgba(11,127,117,.30);
            border-radius: 24px;
            padding: 18px 20px;
            margin: 10px 0 16px 0;
            background: linear-gradient(135deg,#EAF7F1 0%,#FFFFFF 72%,#FFF3C4 100%);
            box-shadow: 0 10px 32px rgba(11,127,117,.10);
            break-inside: avoid; page-break-inside: avoid;
        }
        .tre-report-title-card h1 { margin:0; color:#064E47; font-size:26px; line-height:1.2; }
        .tre-report-title-card p { margin:8px 0 0 0; color:#475569; font-size:13px; }
        .tre-report-section-title {
            margin: 26px 0 10px 0;
            padding: 10px 14px;
            border-left: 7px solid #0B7F75;
            border-radius: 14px;
            background: linear-gradient(90deg,#EAF7F1 0%,#FFFFFF 100%);
            color: #064E47;
            font-size: 19px;
            font-weight: 900;
            break-after: avoid; page-break-after: avoid;
        }
        .tre-report-table-title {
            font-weight: 900; color:#064E47; margin: 16px 0 6px 0; font-size: 14px;
        }
        .tre-report-figure-title {
            font-weight: 900; color:#064E47; margin: 12px 0 4px 0; font-size: 14px;
        }
        .tre-static-table-block {
            margin: 12px 0 18px 0;
            break-inside: auto; page-break-inside: auto;
        }
        .tre-static-table-wrap {
            width: 100%;
            overflow-x: auto;
            border: 1px solid rgba(11,127,117,.18);
            border-radius: 12px;
            background: white;
            box-shadow: 0 4px 18px rgba(15,23,42,.045);
        }
        table.tre-static-table {
            width: 100%;
            border-collapse: collapse;
            font-size: 11px;
            line-height: 1.28;
            color: #0F172A;
        }
        table.tre-static-table th {
            background: linear-gradient(135deg,#D9F0EA 0%,#EAF7F1 100%);
            color: #064E47;
            font-weight: 900;
            border: 1px solid #CBD5E1;
            padding: 7px 8px;
            text-align: left;
            vertical-align: top;
            white-space: normal;
        }
        table.tre-static-table td {
            border: 1px solid #E2E8F0;
            padding: 6px 8px;
            vertical-align: top;
            white-space: pre-wrap;
            overflow-wrap: anywhere;
            word-break: normal;
        }
        table.tre-static-table th.col-code, table.tre-static-table td.col-code {
            min-width: 76px;
            width: 76px;
            max-width: 96px;
            white-space: nowrap !important;
            overflow-wrap: normal !important;
            word-break: keep-all !important;
            text-align: center;
            font-weight: 850;
        }
        table.tre-table-peer th.col-code, table.tre-table-peer td.col-code {
            min-width: 100px;
            width: 100px;
            max-width: 120px;
            white-space: nowrap !important;
        }
        table.tre-static-table th.col-period, table.tre-static-table td.col-period {
            min-width: 92px;
            width: 92px;
            max-width: 116px;
            white-space: nowrap !important;
            overflow-wrap: normal !important;
            word-break: keep-all !important;
            text-align: center;
            font-weight: 850;
        }
        table.tre-table-roic {
            font-size: 9.6px;
            line-height: 1.15;
        }
        table.tre-table-roic th, table.tre-table-roic td {
            padding: 4px 4.5px;
            vertical-align: middle;
        }
        table.tre-table-roic td:not(.col-long-text) {
            white-space: nowrap;
            overflow-wrap: normal;
            word-break: keep-all;
        }
        table.tre-static-table tr:nth-child(even) td { background: #FBFEFC; }
        table.tre-static-table td.num-pos, table.tre-static-table td.heat-green { background: #D1FAE5; color: #064E3B; font-weight: 750; }
        table.tre-static-table td.heat-green-strong { background: #A7F3D0; color: #064E3B; font-weight: 850; }
        table.tre-static-table td.num-neg, table.tre-static-table td.heat-red { background: #FEE2E2; color: #991B1B; font-weight: 750; }
        table.tre-static-table td.heat-yellow { background: #FFF3C4; color: #78350F; font-weight: 750; }
        table.tre-static-table td.heat-orange { background: #FED7AA; color: #7C2D12; font-weight: 750; }
        .tre-table-footnote { color:#64748B; font-size:11px; margin-top:6px; font-style:italic; }
        .tre-table-notes-block {
            border-left: 5px solid #F5B21B;
            background: linear-gradient(135deg,#FFF9DB 0%,#FFFFFF 100%);
            border-radius: 14px;
            padding: 10px 12px;
            margin: 6px 0 18px 0;
            break-inside: auto; page-break-inside: auto;
        }
        .tre-table-notes-title { font-weight: 900; color:#064E47; margin-bottom: 8px; font-size: 12.5px; }
        .tre-table-note-card {
            background: rgba(255,255,255,.82);
            border: 1px solid rgba(245,178,27,.35);
            border-radius: 10px;
            padding: 7px 9px;
            margin: 6px 0;
            font-size: 11.3px;
            line-height: 1.42;
            color: #1F2937;
            break-inside: avoid; page-break-inside: avoid;
        }
        .tre-table-note-card.muted { color:#64748B; font-style: italic; }
        @media print {
          .tre-print-toolbar, .tre-report-export-box { display: none !important; }
          .tre-report-title-card h1 { font-size: 22px !important; }
          .tre-report-section-title { font-size: 16px !important; margin-top: 18px !important; }
          .tre-static-table-wrap { overflow: visible !important; border-radius: 8px !important; box-shadow: none !important; }
          table.tre-static-table { font-size: 8.4px !important; line-height: 1.16 !important; table-layout: auto !important; }
          table.tre-static-table th, table.tre-static-table td { padding: 3.6px 4.4px !important; }
          table.tre-static-table th.col-code, table.tre-static-table td.col-code { min-width: 68px !important; width: 68px !important; white-space: nowrap !important; }
          table.tre-table-peer th.col-code, table.tre-table-peer td.col-code { min-width: 92px !important; width: 92px !important; }
          table.tre-static-table th.col-period, table.tre-static-table td.col-period { min-width: 84px !important; width: 84px !important; max-width: 100px !important; white-space: nowrap !important; word-break: keep-all !important; overflow-wrap: normal !important; }
          table.tre-table-roic { font-size: 6.8px !important; line-height: 1.08 !important; }
          table.tre-table-roic th, table.tre-table-roic td { padding: 2.2px 2.6px !important; }
          table.tre-table-roic td:not(.col-long-text) { white-space: nowrap !important; overflow-wrap: normal !important; word-break: keep-all !important; }
          .tre-table-note-card { font-size: 8.5px !important; padding: 4px 5px !important; margin: 3px 0 !important; }
          .tre-table-notes-title { font-size: 9.5px !important; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        f"""
        <div class="tre-report-title-card">
          <h1>📄 {_escape_html_text(package.title)}</h1>
          <p>Báo cáo tổng hợp toàn bộ nội dung: Tổng quan doanh nghiệp · Định giá chuyên sâu · So sánh doanh nghiệp nếu đã chạy.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if show_export_hint:
        st.markdown('<div class="tre-print-toolbar">', unsafe_allow_html=True)
        _render_browser_print_pdf_button(
            widget_key=f"consolidated_full_report_{_safe_ticker(package.ticker)}",
            button_text="🖨️ In Báo cáo tổng hợp toàn bộ nội dung / Save as PDF",
            help_text="Trang này đã render toàn bộ nội dung Tổng quan doanh nghiệp + Định giá chuyên sâu + So sánh doanh nghiệp. Khi hộp thoại in mở ra, chọn Save as PDF hoặc Microsoft Print to PDF, A4 ngang, bật Background graphics để giữ màu giống app.",
        )
        st.markdown('</div>', unsafe_allow_html=True)

    for section in package.sections:
        st.markdown(f"<div class='tre-report-section-title'>{_escape_html_text(section.title)}</div>", unsafe_allow_html=True)
        for idx, paragraph in enumerate(section.paragraphs or []):
            text = str(paragraph or "").strip()
            if not text:
                continue
            tone = "yellow" if any(k in text.lower() for k in ["khuyến nghị", "cảnh báo", "mos", "rủi ro", "chưa đủ"]) else "green" if idx == 0 else "normal"
            st.markdown(_render_report_card_html("Nhận xét / đánh giá" if idx == 0 else "Nhận xét phân tích", text, tone=tone), unsafe_allow_html=True)
        for fig_title, fig in section.figures or []:
            try:
                st.markdown(f"<div class='tre-report-figure-title'>📈 {_escape_html_text(fig_title)}</div>", unsafe_allow_html=True)
                fig.update_layout(height=430, margin=dict(l=32, r=24, t=42, b=32), paper_bgcolor="white", plot_bgcolor="white")
                st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
            except Exception as exc:
                st.warning(f"Không hiển thị được biểu đồ {fig_title}: {exc}")
        for table_title, df in section.tables or []:
            raw_df = _df(df)
            if raw_df.empty:
                continue
            st.markdown(_render_static_print_table_html(raw_df, table_title, max_rows=max_rows_per_table), unsafe_allow_html=True)

    if package.warnings:
        st.markdown("<div class='tre-report-section-title'>⚠️ Cảnh báo xuất báo cáo</div>", unsafe_allow_html=True)
        st.markdown(_render_report_card_html("Cảnh báo", "\n".join([str(w) for w in package.warnings]), tone="yellow"), unsafe_allow_html=True)

def export_full_analysis_report(
    company: CompanyOverview,
    annual_df: pd.DataFrame,
    quarterly_df: pd.DataFrame,
    *,
    export_format_label: str = "Excel (.xlsx)",
    report_dir: str | Path = DEFAULT_REPORT_DIR,
    **kwargs: Any,
) -> tuple[bytes, str, str, Path]:
    package = build_report_package(company, annual_df, quarterly_df, **kwargs)
    ext, mime = EXPORT_FORMATS.get(export_format_label, EXPORT_FORMATS["Excel (.xlsx)"])
    out_name = f"{package.ticker}_FULL_ANALYSIS_REPORT.{ext}"
    out_path = Path(report_dir) / out_name
    if ext == "xlsx":
        data = export_report_xlsx(package, out_path)
    elif ext == "docx":
        data = export_report_docx(package, out_path)
    elif ext == "pdf":
        data = export_report_pdf(package, out_path)
    else:
        raise ValueError(f"Định dạng chưa hỗ trợ: {export_format_label}")
    return data, out_name, mime, out_path


def render_full_report_export_box(
    company: CompanyOverview,
    annual_df: pd.DataFrame,
    quarterly_df: pd.DataFrame,
    *,
    report_dir: str | Path = DEFAULT_REPORT_DIR,
    widget_key: str = "full_report_export",
    expanded: bool = False,
    **kwargs: Any,
) -> None:
    """V23.40 UI: only expose the consolidated full-module report.

    Excel/Word/server-side PDF export paths remain as internal fallback functions, but the app UI no longer
    offers separate export methods. The single user-facing flow is:
    open "Báo cáo tổng hợp toàn bộ nội dung" -> print/Save as PDF, so the PDF keeps app-like cards,
    tables, charts, colors and layout.
    """
    import streamlit as st

    with st.expander("📄 Xuất Báo cáo tổng hợp toàn bộ nội dung", expanded=expanded):
        st.markdown(
            "App hiện chỉ dùng **một luồng xuất báo cáo**: mở trang **Báo cáo tổng hợp toàn bộ nội dung** rồi "
            "bấm **In / Save as PDF**. Trang này sẽ tổng hợp toàn bộ từng tab của từng phần, bao gồm "
            "đánh giá, chấm điểm, phân tích, biểu đồ và bảng số liệu chính, để giữ format giống app nhất."
        )
        try:
            st.page_link(
                "pages/04_Bao_cao_tong_hop.py",
                label="📄 Mở Báo cáo tổng hợp toàn bộ nội dung",
                icon="📄",
            )
        except Exception:
            st.info("Hãy chọn trang 'Báo cáo tổng hợp toàn bộ nội dung' ở sidebar để in/Save as PDF.")
        st.caption("Khi in PDF: chọn A4 ngang/Landscape và bật Background graphics để giữ màu card, bảng và biểu đồ.")
