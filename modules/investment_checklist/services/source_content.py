from __future__ import annotations

"""Append-only source text used by the governed AI research workflow.

The binary upload is never written to the database.  The app stores only normalized,
locator-marked text plus a SHA-256 fingerprint so citations can be verified and audited.
"""

from io import BytesIO
import hashlib
from pathlib import Path
import re
from typing import Any

from ..repositories.sqlite_repository import ValidationError


MAX_UPLOAD_BYTES = 25 * 1024 * 1024
MAX_SOURCE_CONTENT_CHARS = 2_000_000
MIN_SOURCE_CONTENT_CHARS = 40
SUPPORTED_SOURCE_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".csv", ".json")


def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\x00", "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def parse_page_selection(value: str, total_pages: int) -> list[int]:
    """Return zero-based PDF page indexes from a human page expression."""
    if total_pages < 1:
        return []
    raw = str(value or "").strip()
    if not raw:
        return list(range(total_pages))
    pages: set[int] = set()
    for token in raw.split(","):
        token = token.strip()
        if not token:
            continue
        if "-" in token:
            start_text, end_text = token.split("-", 1)
            try:
                start, end = int(start_text), int(end_text)
            except ValueError as exc:
                raise ValidationError("Phạm vi trang PDF phải có dạng 1-5,8,12-14.") from exc
            if start > end:
                raise ValidationError("Trang bắt đầu không được lớn hơn trang kết thúc.")
        else:
            try:
                start = end = int(token)
            except ValueError as exc:
                raise ValidationError("Phạm vi trang PDF phải có dạng 1-5,8,12-14.") from exc
        if start < 1 or end > total_pages:
            raise ValidationError(f"Trang PDF phải nằm trong khoảng 1–{total_pages}.")
        pages.update(range(start - 1, end))
    if not pages:
        raise ValidationError("Phạm vi trang PDF không được để trống.")
    return sorted(pages)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1258", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValidationError("Không thể giải mã tài liệu văn bản.")


def _extract_pdf(data: bytes, page_selection: str) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - deployment dependency guard
        raise ValidationError("Thiếu thư viện pypdf để đọc PDF.") from exc
    try:
        reader = PdfReader(BytesIO(data))
    except Exception as exc:
        raise ValidationError(f"Không thể mở PDF: {exc}") from exc
    page_indexes = parse_page_selection(page_selection, len(reader.pages))
    sections: list[str] = []
    for index in page_indexes:
        try:
            text = _clean_text(reader.pages[index].extract_text() or "")
        except Exception as exc:
            raise ValidationError(f"Không thể đọc trang {index + 1} của PDF: {exc}") from exc
        sections.append(f"[[PAGE {index + 1}]]\n{text or '[NO EXTRACTABLE TEXT]'}")
    scope = f"PDF pages {','.join(str(index + 1) for index in page_indexes)}"
    return "\n\n".join(sections), scope


def _extract_docx(data: bytes) -> tuple[str, str]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - deployment dependency guard
        raise ValidationError("Thiếu thư viện python-docx để đọc DOCX.") from exc
    try:
        document = Document(BytesIO(data))
    except Exception as exc:
        raise ValidationError(f"Không thể mở DOCX: {exc}") from exc
    sections: list[str] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        text = _clean_text(paragraph.text)
        if text:
            sections.append(f"[[PARAGRAPH {index}]]\n{text}")
    for table_no, table in enumerate(document.tables, 1):
        for row_no, row in enumerate(table.rows, 1):
            cells = [_clean_text(cell.text) for cell in row.cells]
            if any(cells):
                sections.append(f"[[TABLE {table_no} ROW {row_no}]]\n" + " | ".join(cells))
    return "\n\n".join(sections), "DOCX paragraph/table markers"


def _extract_plain_text(data: bytes) -> tuple[str, str]:
    raw = _clean_text(_decode_text(data))
    sections = [f"[[LINE {index}]] {line}" for index, line in enumerate(raw.split("\n"), 1) if line.strip()]
    return "\n".join(sections), "Text line markers"


def extract_document_text(filename: str, data: bytes, *, pdf_pages: str = "") -> dict[str, Any]:
    safe_name = Path(str(filename or "source.txt")).name
    suffix = Path(safe_name).suffix.lower()
    if suffix not in SUPPORTED_SOURCE_EXTENSIONS:
        raise ValidationError("Chỉ hỗ trợ PDF, DOCX, TXT, MD, CSV và JSON.")
    if not data:
        raise ValidationError("Tệp nguồn đang trống.")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValidationError(f"Tệp nguồn không được vượt quá {MAX_UPLOAD_BYTES // (1024 * 1024)} MB.")

    if suffix == ".pdf":
        content, scope = _extract_pdf(data, pdf_pages)
        locator_scheme = "page"
        content_type = "application/pdf"
    elif suffix == ".docx":
        content, scope = _extract_docx(data)
        locator_scheme = "paragraph_table"
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        content, scope = _extract_plain_text(data)
        locator_scheme = "line"
        content_type = {
            ".md": "text/markdown", ".csv": "text/csv", ".json": "application/json",
        }.get(suffix, "text/plain")

    content = _clean_text(content)
    if len(content) < MIN_SOURCE_CONTENT_CHARS:
        raise ValidationError("Tài liệu không có đủ nội dung văn bản để AI phân tích.")
    if len(content) > MAX_SOURCE_CONTENT_CHARS:
        raise ValidationError(
            f"Nội dung trích xuất vượt {MAX_SOURCE_CONTENT_CHARS:,} ký tự; hãy giới hạn phạm vi trang."
        )
    return {
        "content_text": content,
        "content_type": content_type,
        "locator_scheme": locator_scheme,
        "original_filename": safe_name,
        "scope_label": scope,
        "content_hash": _content_hash(content),
        "char_count": len(content),
    }


def create_source_content_version(
    repo,
    *,
    company_ref_id: int,
    source_id: int,
    content_text: str,
    content_type: str = "text/plain",
    locator_scheme: str = "line",
    original_filename: str = "",
    scope_label: str = "",
    actor: str = "analyst",
) -> int:
    content = _clean_text(content_text)
    if len(content) < MIN_SOURCE_CONTENT_CHARS:
        raise ValidationError("Nội dung nguồn phải có ít nhất 40 ký tự.")
    if len(content) > MAX_SOURCE_CONTENT_CHARS:
        raise ValidationError(f"Nội dung nguồn không được vượt quá {MAX_SOURCE_CONTENT_CHARS:,} ký tự.")
    content_type = str(content_type or "text/plain").strip()[:200]
    locator_scheme = str(locator_scheme or "line").strip()[:100]
    original_filename = Path(str(original_filename or "")).name[:500]
    scope_label = str(scope_label or "").strip()[:2000]
    actor = str(actor or "analyst").strip()[:200]
    fingerprint = _content_hash(content)

    with repo._conn() as c:
        source = repo._d(c.execute("SELECT * FROM research_sources WHERE id=?", (int(source_id),)).fetchone())
        if not source or int(source["company_ref_id"]) != int(company_ref_id):
            raise ValidationError("Nguồn không thuộc doanh nghiệp đang phân tích.")
        if source["status"] != "active":
            raise ValidationError("Nguồn đã archived; không thể thêm content version.")
        duplicate = c.execute(
            "SELECT id FROM research_source_contents WHERE source_id=? AND content_hash=?",
            (int(source_id), fingerprint),
        ).fetchone()
        if duplicate:
            raise ValidationError(f"Nội dung này đã tồn tại (Content #{duplicate['id']}).")
        row = c.execute(
            "SELECT COALESCE(MAX(version_no),0)+1 next_version FROM research_source_contents WHERE source_id=?",
            (int(source_id),),
        ).fetchone()
        fields = {
            "company_ref_id": int(company_ref_id),
            "source_id": int(source_id),
            "version_no": int(row["next_version"]),
            "content_type": content_type,
            "locator_scheme": locator_scheme,
            "original_filename": original_filename or None,
            "scope_label": scope_label or None,
            "content_text": content,
            "content_hash": fingerprint,
            "char_count": len(content),
            "created_by": actor,
        }
        cur = c.execute(
            f"INSERT INTO research_source_contents({','.join(fields)}) VALUES({','.join('?' for _ in fields)})",
            tuple(fields.values()),
        )
        content_id = int(cur.lastrowid)
        audit_after = {key: value for key, value in fields.items() if key != "content_text"}
        audit_after["id"] = content_id
        repo._audit(
            c, company_ref_id=company_ref_id, actor=actor, action="create_version",
            entity_type="research_source_content", entity_id=content_id, after=audit_after,
        )
        return content_id


def list_source_contents(repo, company_ref_id: int, *, source_id: int | None = None) -> list[dict[str, Any]]:
    sql = """SELECT c.id,c.company_ref_id,c.source_id,c.version_no,c.content_type,c.locator_scheme,
    c.original_filename,c.scope_label,c.content_hash,c.char_count,c.created_by,c.created_at,
    s.title source_title,s.status source_status
    FROM research_source_contents c JOIN research_sources s ON s.id=c.source_id
    WHERE c.company_ref_id=?"""
    params: list[Any] = [int(company_ref_id)]
    if source_id is not None:
        sql += " AND c.source_id=?"
        params.append(int(source_id))
    sql += " ORDER BY c.source_id,c.version_no DESC,c.id DESC"
    with repo._conn() as c:
        return [dict(row) for row in c.execute(sql, tuple(params))]


def get_source_content(repo, content_id: int, *, conn=None) -> dict[str, Any] | None:
    sql = """SELECT c.*,s.title source_title,s.status source_status,s.source_hash
    FROM research_source_contents c JOIN research_sources s ON s.id=c.source_id WHERE c.id=?"""
    if conn is not None:
        return repo._d(conn.execute(sql, (int(content_id),)).fetchone())
    with repo._conn() as c:
        return repo._d(c.execute(sql, (int(content_id),)).fetchone())


def latest_source_content(repo, source_id: int, *, conn=None) -> dict[str, Any] | None:
    sql = """SELECT c.*,s.title source_title,s.status source_status,s.source_hash
    FROM research_source_contents c JOIN research_sources s ON s.id=c.source_id
    WHERE c.source_id=? ORDER BY c.version_no DESC,c.id DESC LIMIT 1"""
    if conn is not None:
        return repo._d(conn.execute(sql, (int(source_id),)).fetchone())
    with repo._conn() as c:
        return repo._d(c.execute(sql, (int(source_id),)).fetchone())


__all__ = [
    "MAX_SOURCE_CONTENT_CHARS", "MAX_UPLOAD_BYTES", "SUPPORTED_SOURCE_EXTENSIONS",
    "create_source_content_version", "extract_document_text", "get_source_content",
    "latest_source_content", "list_source_contents", "parse_page_selection",
]
